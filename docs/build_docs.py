"""
Bygger Teknisk dokumentation.docx og Brugervejledning.docx til Lønsystemet.
Kør med: python build_docs.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = Path(__file__).parent

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x6B)
MID_BLUE  = RGBColor(0x2E, 0x6D, 0xB8)
LIGHT_BLUE = RGBColor(0xD5, 0xE8, 0xF5)
GRAY      = RGBColor(0x60, 0x60, 0x60)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)


# ─── helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = "Arial"
    return run


def heading(doc, text, level=1, num=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    label = f"{num}  " if num else ""
    run = p.add_run(label + text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE if level == 2 else BLACK
    return p


def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = "Arial"
        rb.font.size = Pt(11)
        r = p.add_run(text)
    else:
        r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    return p


def cover(doc, title, subtitle, date_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Arial"
    r.font.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Arial"
    r2.font.size = Pt(14)
    r2.font.color.rgb = MID_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(60)
    r3 = p3.add_run("Poul Schou A/S  ·  CVR 13246505")
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(date_str)
    r4.font.name = "Arial"
    r4.font.size = Pt(11)
    r4.font.color.rgb = GRAY

    doc.add_page_break()


def two_col_table(doc, rows, col1_w=Cm(5), col2_w=Cm(11.5)):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for label, val in rows:
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, "D5E8F5")
        set_cell_borders(c0, "9BBDD6")
        set_cell_borders(c1, "9BBDD6")
        c0.width = col1_w
        c1.width = col2_w
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True; r0.font.name = "Arial"; r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = "Arial"; r1.font.size = Pt(10)
    doc.add_paragraph()


def header_table(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "1B3A6B")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.color.rgb = WHITE
    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Arial"; r.font.size = Pt(10)
    doc.add_paragraph()


def note_box(doc, text, label="OBS"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    rb = p.add_run(f"{label}: ")
    rb.bold = True; rb.font.name = "Arial"; rb.font.size = Pt(11)
    rb.font.color.rgb = MID_BLUE
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  TEKNISK DOKUMENTATION
# ══════════════════════════════════════════════════════════════════════════════

def build_teknisk():
    doc = Document()

    # Sider og margener
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)

    # Forside
    cover(doc, "Lønsystem", "Teknisk dokumentation", "Juni 2026")

    # ── 1. Systemarkitektur ────────────────────────────────────────────────
    heading(doc, "Systemarkitektur", 1, "1")
    body(doc, (
        "Lønsystemet er en webbaseret applikation til registrering og lønberegning for "
        "lastbilchauffører hos Poul Schou A/S. Systemet kører på én central server og "
        "tilgås via en webbrowser på det lokale netværk."
    ))

    heading(doc, "Teknologistack", 2, "1.1")
    header_table(doc,
        ["Komponent", "Teknologi", "Formål"],
        [
            ["Backend", "Python 3.13 + FastAPI", "REST-API, forretningslogik, filgenerering"],
            ["Database", "SQLite (WAL-mode)", "Vedvarende lagring af medarbejder- og aktivitetsdata"],
            ["Frontend", "HTML5 + Vanilla JavaScript + CSS", "Brugerflade – ét HTML-dokument"],
            ["Server", "Uvicorn (ASGI)", "Serverer FastAPI-applikationen"],
            ["PDF-generering", "ReportLab", "Timesedler til medarbejdere"],
            ["Excel-generering", "openpyxl", "Prøvekørsel og satser"],
            ["Filparsing", "Binær parsing (ingen eksternt bibliotek)", "Indlæsning af .ddd tachografdata"],
        ],
        [Cm(3), Cm(5), Cm(8.5)]
    )

    heading(doc, "Mappestruktur", 2, "1.2")
    body(doc, "Projektmappen er organiseret i fire hoveddele:")
    two_col_table(doc, [
        ["app/",          "Al Python-kode, statiske filer, database og Excel-satser"],
        ["docs/",         "Projektdokumentation (.md-filer og denne fil)"],
        ["præsentation/", "Præsentationsbygger og slides (til intern brug)"],
        ["Rodmappen",     "Brugerdokumenter: overenskomst-PDF, kravdokument, statusmøde-pptx m.m."],
    ])

    heading(doc, "Centrale filer i app/", 2, "1.3")
    two_col_table(doc, [
        ["main.py",                   "Applikationens indgangspunkt. Registrerer alle routere, monterer statiske filer og serverer index.html med cache-busting."],
        ["routers/activities.py",     "API-endepunkter for aktivitetshåndtering (opret, godkend, ret, opdel, fortryd)."],
        ["routers/employees.py",      "API-endepunkter for medarbejderstamdata og overenskomsttyper."],
        ["routers/payroll_router.py", "Lønkørsel: preview, prøvekørsel (Excel), Danløn CSV og PDF-timesedler."],
        ["routers/import_ddd.py",     "Import af .ddd-filer via fildialoger (tkinter) og batch-import."],
        ["calculators/overtime.py",   "Overtids- og tillægsberegning pr. aktivitet."],
        ["calculators/rates_loader.py","Indlæser timesatser og overtidssatser fra DB (Stamdata). Excel-funktioner bruges kun som fallback ved tom DB."],
        ["routers/stamdata.py",       "CRUD-endepunkter for stamdata: overenskomsttyper, overtidssatser, tillæg, løntypekoder og fraværstyper. Kræver 'stamdata'-rettighed."],
        ["calculators/pay_period.py", "Beregner og opretter 14-dages lønperioder."],
        ["database/models.py",        "SQLAlchemy-modeller (Employee, Activity, PayPeriod m.fl.)."],
        ["database/schemas.py",       "Pydantic-skemaer til validering af API-input og -output."],
        ["parsers/ddd_parser.py",     "Binær parser for EU-tachografdata (.ddd-format)."],
        ["templates/index.html",      "Hele frontend-applikationen (ét HTML-dokument)."],
        ["static/js/app.js",          "JavaScript-applikationslogik (~1700 linjer)."],
    ])

    # ── 2. Databasemodel ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Databasemodel", 1, "2")
    body(doc, (
        "Databasen er SQLite og oprettes automatisk ved første opstart. WAL-mode "
        "(Write-Ahead Logging) aktiveres, så flere samtidige læsere kan tilgå databasen "
        "mens en skrivning foregår."
    ))

    heading(doc, "Employee – medarbejdere", 2, "2.1")
    header_table(doc,
        ["Felt", "Type", "Beskrivelse"],
        [
            ["employee_number",       "String (unik)",   "Lønnummer – primær identifikator i lønsystemet"],
            ["tachograph_card_number","String (unik)",   "EU-førerkortnummer – bruges til at matche .ddd-filer"],
            ["first_name / last_name","String",          "Navn"],
            ["address / postal_code", "String (opt.)",   "Adresse"],
            ["email / phone / mobile","String (opt.)",   "Kontaktoplysninger"],
            ["agreement_kind",        "String",          "Nøgle fra master_agreement_kinds.key (se afsnit 2.5) – ikke længere en fast enum-kolonne. Systemnøglerne hourly_fixed/hourly_flexible styrer stadig overtidsberegningen (se afsnit 5.5); nye, brugeroprettede nøgler springes automatisk over i den"],
            ["agreement_type",        "String",          "Overenskomsttype fra Excel-filen – bestemmer timesats"],
            ["fuldloennet",           "Boolean",         "Fuldlønnet status"],
            ["active",                "Boolean",         "Aktiv medarbejder (filtreres ved opslag)"],
            ["hire_date",             "Date",            "Ansættelsesdato"],
            ["termination_date",      "Date",            "Fratrædelsesdato (default 31-12-9999)"],
            ["work_schedule",         "JSON",            '{"even":[man..søn], "odd":[man..søn]} – normaltimer pr. dag i lig/ulige uge'],
            ["anciennitet_dismissed_at","DateTime (opt.)","Tidspunkt for afvisning af anciennitetsvarsel (nulstilles ved overenskomstskifte)"],
        ]
    )
    body(doc, (
        "Disponentgrupper (afdelinger) er ikke et enkelt felt på medarbejderen, men en mange-til-mange-relation "
        "via tabellerne dispatcher_groups (navn, beskrivelse) og employee_dispatcher_groups (kobling). En medarbejder "
        "kan tilhøre 0-N grupper samtidig, og der er ingen 'primær' gruppe. Grupperne administreres under "
        "Stamdata → Disponentgrupper (opret/omdøb/slet); en medarbejder tilknyttes sine grupper via afkrydsningsbokse "
        "i medarbejder-modalen. Ved filtrering (aktivitetstabel, fraværsoversigt-eksport) vises en medarbejder under "
        "alle sine tilknyttede grupper. Frem til 27. juli 2026 lå dette som en enkelt tekststreng på medarbejderen – "
        "denne kolonne er fjernet og data migreret automatisk til de nye tabeller ved første opstart efter opgraderingen."
    ))

    heading(doc, "Activity – aktiviteter", 2, "2.2")
    header_table(doc,
        ["Felt", "Type", "Beskrivelse"],
        [
            ["employee_id",           "FK → Employee",  "Tilknyttet medarbejder"],
            ["pay_period_id",         "FK → PayPeriod", "Tilknyttet lønperiode"],
            ["trip_number",           "String (opt.)",  "Turnummer (max 6 tegn)"],
            ["source",                "Enum",           "tachograph (DDD-import) / manual (manuelt oprettet)"],
            ["activity_type",         "String(50)",     "normal / ferie / fri / afspadsering / skole_kursus / overnatning"],
            ["start_time / end_time", "DateTime",       "Start- og sluttidspunkt"],
            ["availability_time_pct", "Decimal",        "Rådighedstid i % (fra tachograf)"],
            ["rest_pause_pct",        "Decimal",        "Hvil/pause i % (fra tachograf)"],
            ["other_work_pct",        "Decimal",        "Andet arbejde i % (fra tachograf)"],
            ["driving_pct",           "Decimal",        "Kørsel i % (fra tachograf)"],
            ["loading_minutes",       "Integer (opt.)", "Pålæsningstid i minutter"],
            ["unloading_minutes",     "Integer (opt.)", "Aflæsningstid i minutter"],
            ["pause_intervals",       "JSON",           '[ ["ISO-start","ISO-slut"], ... ] – pauseintervaller'],
            ["segments",              "JSON",           '[ ["ISO-start","ISO-slut","type"], ... ] – tachografsegmenter'],
            ["original_start_time",   "DateTime (opt.)","Original tachograftid inden første rettelse"],
            ["original_end_time",     "DateTime (opt.)","Original tachograftid inden første rettelse"],
            ["status",                "Enum",           "pending / approved / deactivated"],
            ["approved_by",           "String (opt.)",  "Initialer på godkender (kun ved godkendelse)"],
            ["approved_at",           "DateTime (opt.)","Godkendelsestidspunkt"],
            ["deactivated_by",        "String (opt.)",  "Initialer på den der deaktiverede (kun ved deaktivering)"],
            ["comment",               "Text (opt.)",    "Kommentar (påkrævet ved godkendelse < 4 timer)"],
            ["parent_activity_id",    "FK → Activity",  "Peger på moderaktivitet ved opdeling"],
            ["split_part",            "Integer (opt.)", "1 eller 2 – rækkefølge af opdelingsdele"],
        ]
    )

    heading(doc, "PayPeriod – lønperioder", 2, "2.3")
    two_col_table(doc, [
        ["start_date / end_date", "Altid mandag–søndag, 14 dage"],
        ["status",                "open (aktiv) / preview (under kørsel) / closed (afsluttet)"],
        ["closed_at / closed_by", "Tidspunkt og bruger ved afslutning"],
    ])

    heading(doc, "PayrollRun – lønkørsler", 2, "2.4")
    two_col_table(doc, [
        ["pay_period_id", "FK til lønperioden"],
        ["run_type",      '"preview" (prøvekørsel) eller "final" (endelig kørsel)'],
        ["run_at",        "Tidspunkt for kørslen"],
        ["run_by",        "Initialer på den der kørte løn"],
        ["csv_path",      "Sti til genereret Danløn CSV-fil"],
        ["excel_path",    "Sti til genereret Excel-fil"],
    ])

    heading(doc, "Stamdata – masterdatatabeller", 2, "2.5")
    body(doc, (
        "Seks tabeller holder systemets masterdata. De fleste seedes automatisk fra Excel-filerne "
        "ved første opstart (Excel-filerne bruges herefter kun som fallback hvis en tabel er tom); "
        "master_agreement_kinds seedes i stedet fra to faste, hardkodede systemrækker (se nedenfor). "
        "Alle tabeller redigeres derefter via Stamdata-modulet i systemets brugerflade."
    ))
    header_table(doc,
        ["Tabel", "Indhold", "Vigtige felter"],
        [
            ["master_agreement_types",  "Overenskomsttyper og timesatser", "name (unik), hourly_rate"],
            ["master_agreement_kinds",  "Aftaletyper (\"Aftale\"-feltet på medarbejderen)", "key (unik, fast efter oprettelse), label, is_active, is_user_created, requires_agreement_type, sort_order"],
            ["master_overtime_rates",   "De tre overtidssatser",           "label (unik), rate"],
            ["master_supplement_rates", "Salttillæg, Overnatning, Dagpenge §56", "label (unik), rate"],
            ["master_pay_types",        "Løntypekoder til Danløn CSV",     "code_key (unik), danloen_code, include_in_csv, csv_quantity_type, csv_rate_source, csv_include_rate, csv_include_total, sort_order"],
            ["master_absence_types",    "Fraværstyper der vises i UI",     "label, normalized_key (unik), is_active, is_user_created, sort_order"],
        ]
    )
    body(doc, (
        "Normaliserede nøgler i master_absence_types genereres én gang ved oprettelse og ændres "
        "ikke efterfølgende – de bruges internt i beregnings- og aktivitetslogikken. "
        "Alle fraværstyper kan slettes via Stamdata-UI'et."
    ))
    body(doc, (
        "master_agreement_kinds seedes ved opstart med de to systemrækker hourly_fixed og "
        "hourly_flexible (is_user_created=False) – de kan ikke slettes, men deres label kan redigeres "
        "frit. key sættes kun ved oprettelse og kan ALDRIG ændres bagefter (heller ikke for "
        "brugeroprettede rækker), fordi overtidsberegningen i overtime.py/payroll_router.py grener "
        "direkte på strengværdierne hourly_fixed/hourly_flexible – se afsnit 5.5. Feltet "
        "requires_agreement_type styrer om Overenskomsttype er påkrævet for medarbejdere med den "
        "pågældende aftaletype (håndhæves i routers/employees.py)."
    ))

    heading(doc, "Holidays – helligdagskalender", 2, "2.6")
    body(doc, (
        "Holidays-tabellen gemmer danske helligdage og administreres separat fra de Excel-seedede "
        "masterdatatabeller. Den seedes automatisk ved serveropstart via _seed_holidays() i session.py."
    ))
    header_table(doc,
        ["Felt", "Type", "Beskrivelse"],
        [
            ["date",              "DATE UNIQUE NOT NULL", "Helligdagens dato – UNIQUE constraint forhindrer dubletter"],
            ["name",              "String NOT NULL",      "Navn, fx 'Påskedag', '1. maj', 'Grundlovsdag'"],
            ["half_day_from",     "String(5) NULL",       "'12:00' = fri fra middag; NULL = heldagshelligdag"],
            ["is_auto_generated", "Boolean DEFAULT TRUE", "TRUE = genereret af systemet; FALSE = manuelt oprettet"],
        ]
    )
    for item in [
        "5 løbende år genereres ved opstart (indeværende år + 4 fremover).",
        "Seeding er idempotent – eksisterende datoer springes over.",
        "Påskedato beregnes via anonym Gregoriansk Computus-algoritme i app/calculators/holidays.py.",
        "Deduplicering: faste helligdage vinder over bevægelige ved datokollision "
        "(fx 2028: 2. pinsedag og Grundlovsdag falder begge den 5. juni – Grundlovsdag bevares).",
        "Store Bededag medtages ikke (afskaffet fra 2024).",
    ]:
        bullet(doc, item)
    note_box(doc,
        "half_day_from-feltet er forberedt til brug i fremtidig lønberegning (helligdagstillæg). "
        "Feltet bruges allerede i aktivitetskalenderens '½ fra HH:MM'-badge.",
        "TEKNISK NOTE"
    )

    # ── 3. Lønperioder ────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Lønperioder", 1, "3")
    body(doc, (
        "Lønperioderne er faste 14-dages intervaller fra mandag til søndag. "
        "Systemet bruger en ankerperiode til at beregne alle andre perioder."
    ))
    two_col_table(doc, [
        ["Ankerperiode",  "Mandag 1. juni 2026 – søndag 14. juni 2026"],
        ["Periode-start", "Altid en mandag"],
        ["Periode-slut",  "Altid en søndag (14 dage efter start)"],
        ["Beregning",     "period_start_for_date(d): finder nærmeste mandag via modulo-14 fra ankerpunktet"],
    ])
    body(doc, (
        "get_billing_period() i pay_period.py bestemmer hvilken periode en NY eller RETTET aktivitet skal "
        "placeres i: hører datoen til en periode der allerede har status 'closed' (der er kørt løn), "
        "returneres i stedet den efterfølgende periode (get_or_create_period_for_date() kaldt på "
        "periodens slutdato + 1 dag). Bruges ved oprettelse af aktivitet, rettelse af starttid og "
        "DDD-import. Genåbning af en allerede placeret aktivitet ændrer bevidst ikke dens pay_period_id – "
        "kun oversigtsvisningen (se afsnit 4.4 i Brugervejledningen) sørger for at den stadig ses korrekt."
    ))
    body(doc, (
        "Timefordelingen på medarbejderen skelner mellem 'lige' og 'ulige' uger. "
        "Dette bestemmes af is_even_week(d): ISO-ugenummer % 2 == 0 → 'even' (lige uge). "
        "Mandag 1. juni 2026 er i ISO-uge 23 (ulige), så 'even'-skemaet bruges i ISO-uge 24, 26 osv."
    ))

    # ── 4. DDD-filparsing ─────────────────────────────────────────────────
    heading(doc, "DDD-filparsing (tachografdata)", 1, "4")
    body(doc, (
        "EU-tachografer gemmer data i .ddd-filer (EU Forordning 165/2014). Filen er binær "
        "og indeholder daglige poster med tidsstempler og aktivitetsskift."
    ))

    heading(doc, "Parsingsflow", 2, "4.1")
    for i, step in enumerate([
        ("Filidentifikation", "Kortnummeret matches som landekode (2 bogstaver) + 14 cifre (fx DK00000012666013), men kun de første 14 tegn (landekode + 12 cifre = driverIdentification) er det stabile kortnummer der bruges til medarbejder-matching. De sidste 2 cifre er cardReplacementIndex + cardRenewalIndex og ændrer sig når kortet udskiftes/fornys."),
        ("Lokalisering af poster", "Heuristisk søgning efter timestamp-mønstre identificerer starten på daglige poster."),
        ("Afkodning af ActivityChangeInfo", "2-byte records afkodes bit-for-bit: slot (bit 15-14), aktivitetstype (bit 13-11), minutter fra midnat (bit 10-0). Dato og minutter er UTC."),
        ("Aktivitetstyper", "Rest (hvil), Availability (rådighedstid), Work (andet arbejde), Driving (kørsel)."),
        ("Dagsstart", "Hver dags aktivitetsarray starter altid med en hvil-post ved minut 0 (videreført status fra dagen før, ikke en reel pause). Findes der en ekstra hvil-post lige derefter, er det chaufførens faktiske dagsstart – en kort pause inden arbejdet begynder – og den bruges som visningsstart for dagen."),
        ("Tidszonekonvertering", "Start/sluttid, segmenter og pauseintervaller konverteres fra UTC til dansk lokal tid (Europe/Copenhagen, DST-korrekt via Python-modulet zoneinfo) inden de gemmes."),
        ("Bygning af ParsedActivity", "Start/sluttid, procentfordelinger, pauseintervaller og segmenter samles til et ParsedActivity-objekt."),
        ("Duplikat-tjek/udvidelse", "Findes der allerede en aktivitet med samme medarbejder + starttid, opdateres km-data hvis de mangler. Har den nye fil desuden et SENERE sluttidspunkt end den eksisterende aktivitet (fx fordi kortet først blev læst af midt på en vagt og senere igen efter vagtens afslutning), udvides aktiviteten (sluttid, segmenter, pauser, procentfordeling) i stedet for at blive sprunget over – ellers ville den ekstra tid gå tabt for altid. Var aktiviteten allerede godkendt/deaktiveret, genåbnes den til 'afventende', så den udvidede tid skal godkendes igen (rettet 2026-07-02)."),
        ("Import af aktivitet", "_import_activity() returnerer 'new' (ny aktivitet), 'updated' (eksisterende aktivitet fik km-data udfyldt og/eller blev udvidet med en senere sluttid), 'skipped_unknown_card' (intet førerkortnummer matcher) eller 'skipped_duplicate' (allerede importeret, intet nyt at tilføje) – hver årsag tælles separat."),
        ("Km-start/km-slut", "_extract_daily_odometer() finder et separat array af (km, tidsstempel)-par i filen ved kæde-validering (mindst 5 elementer med præcis 20 bytes' afstand, ingen fast offset). km_start = km-standen tættest på dagens beregnede startminut; km_end = km_start + dagens egen kørte distance."),
    ], 1):
        bullet(doc, f"{step[1]}", f"{i}. {step[0]}: ")

    note_box(doc,
        "zoneinfo kræver på Windows Python-pakken tzdata (installeret via requirements.txt), "
        "da Windows ikke leverer sin egen IANA-tidszonedatabase.",
        "TEKNISK NOTE"
    )
    note_box(doc,
        "En indledende kort pause (fx 1-11 minutter) tælles med i den viste arbejdstid og "
        "fremgår af pause_intervals, men er stadig ubetalt – pause_intervals fratrækkes altid "
        "i lønberegningen uanset hvor i dagen de ligger.",
        "GODT AT VIDE"
    )
    note_box(doc,
        "Førerkort gemmer kun et begrænset antal køretøjsbrug-poster, så km-tabellen dækker "
        "ikke nødvendigvis hele kortets historik. Ældre dage vil derfor mangle km-start/km-slut "
        "– det er en begrænsning i kortets egne data, ikke en fejl i importen.",
        "GODT AT VIDE"
    )

    heading(doc, "Import-flow", 2, "4.2")
    body(doc, (
        "Brugeren klikker 'Vælg filer' eller 'Vælg mappe' i browsergrænsefladen. "
        "En tkinter-dialog (Windows-nativ) åbnes via GET /api/browse-ddd-files eller "
        "/api/browse-ddd-folder. Den valgte sti sendes med POST /api/import-ddd-from, "
        "som parser filerne og gemmer nye aktiviteter i databasen."
    ))
    body(doc, (
        "_process_import_results() i import_ddd.py samler resultatet af alle filer: antal "
        "importeret, opdateret (km-data udfyldt på en eksisterende aktivitet), samt sprunget "
        "over opdelt på årsag (ukendt førerkortnummer – med de konkrete kortnumre – eller "
        "allerede importeret). Filer der giver 0 aktiviteter og eventuelle parse-fejl (også "
        "ved mappe-scanning) indgår i samme resultat. Frontenden viser resultatet i en "
        "pop-up (modal-import-result): 'Importering succesfuld' hvis alt gik igennem uden "
        "sprungne aktiviteter/fejl, ellers en opdelt oversigt over årsagerne."
    ))
    body(doc, (
        "Hver importkørsel logges som én hændelse (action 'ddd_import') i hændelsesloggen "
        "via log_action(), med samme opdelte opsummering i tekstform i details-feltet – så "
        "årsager til sprungne aktiviteter kan slås op bagefter, ikke kun ses i pop-up'en "
        "lige efter kørslen."
    ))

    # ── 5. Overtidsberegning ──────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Overtidsberegning", 1, "5")
    body(doc, (
        "Beregningen udføres i calculators/overtime.py ved funktionen "
        "calculate_overtime(start, end, pause_intervals, normal_hours, rates). "
        "Satserne hentes fra Stamdata-databasen (master_overtime_rates). "
        "Excel bruges kun som fallback hvis DB-tabellen er tom."
    ))

    heading(doc, "De tre tillægstyper", 2, "5.1")
    header_table(doc,
        ["Nøgle", "Tidsvindue", "Betingelse", "Beskrivelse"],
        [
            ["OT_BEFORE\n(1 time før)", "05:00–06:00", "Arbejde i dette interval", "Tillæg for arbejde 1 time inden normal arbejdstid. Sats × timer i vinduet."],
            ["OT_13\n(1-3 timer efter)", "18:00–21:00\n+ 06:00–18:00 over normaltid", "Max 3 timer samlet", "Tillæg for de første 3 overarbejdstimer. Inkluderer dag-timer ud over normaltid (06-18) og aftentimer (18-21)."],
            ["OT_EXTRA\n(Øvrigt overtid)", "21:00–05:00\n+ overarbejde > 3 timer", "Al resterende overarbejde", "Tillæg for natarbejde og overarbejde ud over 3 timer. Højeste sats."],
        ]
    )

    heading(doc, "Beregningsgang", 2, "5.2")
    for step in [
        "Pauseintervaller fratrækkes i de tidsrum de afholdes (korrekt placering i tillægsvindue).",
        "Arbejdstimerne opdeles i 4 vinduer: before (05-06), day (06-18), evening (18-21), night (21-05).",
        "Alle arbejdstimer tæller med i normal_hours (kode 1) uanset vindue – men det REGISTREREDE normaltids-loft (fra medarbejderens timefordeling, fx 7/7,5/8 t) kan kun forbruges af timer i day-vinduet (06-18). Nat-, før- og aftentimer er altid rent tillæg og fortærer ikke loftet (rettet 2026-07-02, se OBS nedenfor).",
        "Day-timer der overskrider det (uforbrugte) normaltids-loft går til OT_13, op til et loft på 3 timer – derefter evening-timer, i den rækkefølge de forekommer kronologisk.",
        "Resterende evening-timer (når OT_13-puljen på 3 timer er brugt op) og alle night-timer går til OT_EXTRA.",
        "Resultat: OvertimeResult med normal_hours, ot_before_hours, ot_13_hours, ot_extra_hours samt beregnede tillæg i kr.",
    ]:
        bullet(doc, step)

    note_box(doc,
        "Pauser skal altid registreres korrekt, da de fratrækkes i det præcise tidsvindue de afholdes. "
        "En pause kl. 20:00 fratrækkes i OT_13-vinduet og reducerer dermed tillægget korrekt.",
        "VIGTIGT"
    )
    note_box(doc,
        "Frem til 2026-07-02 fortærede nat-, før- og aftentimer også normaltids-loftet, hvilket "
        "gjorde for mange senere dagtimer til overarbejde for medarbejdere med tidlig morgenstart. "
        "Reglen er rettet, og calculate_overtime() kan nu tage imod og videresende det resterende "
        "loft (normal_remaining/ot13_remaining) fra en tidligere aktivitet SAMME dag, så loftet "
        "deles pr. dag og ikke nulstilles pr. aktivitet (relevant når en dag er delt i flere "
        "godkendte aktiviteter, fx efter en opdeling).",
        "TEKNISK NOTE"
    )

    heading(doc, "Vagter der krydser midnat", 2, "5.3")
    body(doc, (
        "Normaltids-/OT_13-loftet hører til vagten – den dag den startede – og fortsætter "
        "uændret hen over midnat, så længe det ikke er brugt op. Et eksempel: en vagt fra "
        "fredag kl. 20:30 til lørdag kl. 08:02 bruger fredagens loft for hele vagten; er "
        "fredagens 7 timer slet ikke rørt endnu (fordi aften/nat/'1 time før' aldrig fortærer "
        "det), fylder lørdagsmorgenens dagtimer bare videre op i det samme loft. Dette kræver "
        "ingen særlig kode – calculate_overtime()s kronologiske gennemløb af tidssegmenter "
        "håndterer det automatisk, når vagten IKKE splittes."
    ))
    body(doc, (
        "Søndage og helligdage har derimod en loft-uafhængig regel (se afsnit 5.4): "
        "\"alle kørte timer → kode 9, uanset tidspunkt\". Den regel kan ikke deles med en "
        "efterfølgende hverdags tidsvindues-baserede beregning, så en vagt der STARTER på en "
        "søndag/helligdag splittes altid i to ved midnat – søndagsdelen får søndagens regel, "
        "resten falder tilbage til den følgende dags egne regler."
    ))
    note_box(doc,
        "_split_into_day_pieces() i payroll_router.py udfører splittet. Om en aktivitet skal "
        "splittes afgøres af classify_day(aktivitetens startdato) – kun søndag og helligdage "
        "(inkl. halvdagshelligdage) udløser split. Lørdag og almindelige hverdage splittes ALDRIG, "
        "uanset hvor mange kalenderdage vagten strækker sig over (bekræftet af bruger 2026-07-02).",
        "TEKNISK NOTE"
    )

    heading(doc, "Lørdage, søndage og helligdage (SH-betaling)", 2, "5.4")
    body(doc, (
        "Ud over de tre almindelige tillægstyper gælder særlige regler for lørdage, søndage og "
        "helligdage, implementeret i calculators/day_type.py. Dagtypen afgøres af classify_day() "
        "ud fra ugedag og helligdagskalenderen – en helligdag trumfer altid lørdag/søndag."
    ))
    header_table(doc,
        ["Dagtype", "Regel"],
        [
            ["Lørdag", "Regnes ALTID som en normal hverdag via calculate_overtime(), med lørdagens egne garanterede timer (typisk 0) som loft. Er loftet 0, giver den almindelige dagvindues-logik automatisk 'første op til 3 dagtimer → kode 8, resten → kode 9' – uden særkode (rettet 2026-07-02; den tidligere særregel for lørdag er fjernet)."],
            ["Søndag / heldagshelligdag", "Alle garanterede timer → kode 4 (fuldlønnet) / kode 63 (timelønnet), additivt oveni kørselslønnen. Alle kørte timer → kode 1 + kode 9, UANSET tidspunkt (tids-tillæg tilsidesættes)."],
            ["1. maj (halvdagshelligdag, fri fra 12:00)", "Garanti/2 → kode 4/63. Kørsel før 12:00 → kode 1. Kørsel efter 12:00: første 3 timer → kode 8, resten → kode 9."],
            ["Grundlovsdag (halvdagshelligdag, fri fra 12:00)", "Garanti/2 → kode 4/63. Kørsel før 12:00 → kode 1. Kørsel efter 12:00: ALLE timer → kode 9 (intet kode 8-trin)."],
        ]
    )
    note_box(doc,
        "Reglerne er bekræftet af bruger 2026-06-23 (søndag/helligdag) og 2026-07-02 (lørdag). "
        "Se afsnit 5.3 for hvordan vagter der krydser midnat ind i/ud af disse dage håndteres.",
        "BEMÆRK"
    )

    heading(doc, "Aftale-typer uden fast overtidsregel", 2, "5.5")
    body(doc, (
        "Fra og med indførelsen af Stamdata-fanen \"Aftale\" (se afsnit 2.5 og 6.2) kan der oprettes "
        "flere aftaletyper end de to oprindelige, hourly_fixed og hourly_flexible. Kun disse to "
        "nøgler er kendt af overtidsberegningen ovenfor (afsnit 5.1–5.4)."
    ))
    body(doc, (
        "For en medarbejder hvis agreement_kind IKKE er hourly_fixed eller hourly_flexible, "
        "springer _calculate_employee() (payroll_router.py) hele den daglige/ugentlige "
        "loft-beregning og alle tillæg over. I stedet kaldes calculate_flat_hours() "
        "(calculators/overtime.py), som blot summerer arbejdstiden minus pauser til normal_hours – "
        "ingen ot_before/ot_13/ot_extra, og ingen søndags-/helligdagstillæg (sh_kode8/sh_kode9). "
        "Dette gælder for ALLE dagtyper, ikke kun almindelige hverdage."
    ))
    note_box(doc,
        "Dette er en bevidst, midlertidig afgrænsning: den konkrete beregningslogik for nye "
        "aftaletyper (hvis den overhovedet skal afvige fra flad normaltid) er endnu ikke defineret "
        "og er ikke en del af denne opgave. Se docs/superpowers/specs/2026-08-24-aftale-stamdata-design.md.",
        "TEKNISK NOTE"
    )

    # ── 6. Timesatser og overenskomst ─────────────────────────────────────
    heading(doc, "Timesatser og overenskomst", 1, "6")
    body(doc, (
        "Alle timesatser og overenskomsttyper opbevares i Stamdata-databasen og redigeres "
        "via Stamdata-modulet i systemets brugerflade. Excel-filerne bruges kun til initial "
        "seeding af databasen ved første opstart og er derefter irrelevante for driften."
    ))

    heading(doc, "Stamdata-tabeller og DB-funktioner", 2, "6.1")
    header_table(doc,
        ["Stamdata-tabel", "DB-funktion (rates_loader.py)", "Indhold"],
        [
            ["master_agreement_types",  "load_agreement_types_from_db(db)", "Overenskomstnavn → timesats i kr."],
            ["master_overtime_rates",   "load_overtime_rates_from_db(db)",  "De tre overtidstillæg og deres satser."],
            ["master_supplement_rates", "load_supplement_rates_from_db(db)","Salttillæg, Overnatning, Dagpenge §56."],
            ["master_pay_types",        "load_pay_types_from_db(db)",       "Løntypekoder og Danløn-koder."],
            ["master_absence_types",    "— (forespørges direkte i routers)", "Fraværstyper der vises i UI."],
            ["master_agreement_kinds",  "— (forespørges direkte i routers)", "Aftaletyper til medarbejder-modalens 'Aftale'-dropdown."],
        ]
    )

    body(doc, "")
    bullet(doc, "load_agreement_types_from_db(db): returnerer dict {navn: Decimal sats}", "rates_loader.py: ")
    bullet(doc, "load_overtime_rates_from_db(db): returnerer dict {nøgle: Decimal sats}")
    bullet(doc, "load_salt_supplement_rate_from_db(db): returnerer Decimal sats for salttillæg pr. time")
    bullet(doc, "load_overnight_rate_from_db(db): returnerer Decimal sats for overnatning pr. forekomst")
    bullet(doc, "seniority_variant_exists(type, db): tjekker om der findes en '9 mdr anciennitet'-variant i DB")
    bullet(doc, "Alle _from_db-funktioner har Excel som fallback, der kun aktiveres hvis DB-tabellen er tom")

    body(doc, "")
    body(doc, (
        "Anciennitetsvarsler: systemet kontrollerer løbende om aktive medarbejdere har ≥9 måneder "
        "og en tilsvarende 9-mdr-variant af overenskomsttypen – i så fald vises et varsel i UI. "
        "Varslet styres af tilladelsen 'anciennitet_alert' (checkAnciennitetsAlerts() returnerer tidligt "
        "hvis !state.currentUser.permissions.includes('anciennitet_alert')). "
        "Tilladelsen er som standard slået til for Administrator og Lønbogholder, og kan klikkes til/fra "
        "per rolle i Brugerstyring."
    ))

    heading(doc, "Stamdata-modulet i brugerfladen", 2, "6.2")
    body(doc, (
        "Stamdata-menupunktet (⚙️ Stamdata) i venstre menu giver administratorer adgang til "
        "ni faner med CRUD-funktionalitet for alle masterdatatabeller:"
    ))
    header_table(doc,
        ["Fane", "Indhold", "CRUD-muligheder"],
        [
            ["Overenskomsttyper", "master_agreement_types", "Opret, rediger navn/sats, slet"],
            ["Overtidssatser",    "master_overtime_rates",  "Rediger satser for de tre tillægstyper"],
            ["Tillæg",            "master_supplement_rates","Rediger satser for salttillæg, overnatning, §56"],
            ["Løntypekoder",      "master_pay_types",       "Opret nye, rediger type/kode/CSV-flag/antal-type/sats-kilde/inkl. sats/inkl. total, slet alle"],
            ["Fraværstyper",      "master_absence_types",   "Opret nye, aktiver/deaktiver, slet alle"],
            ["CVR nummer",        "master_cvr_numbers",     "Opret, rediger, sæt standard, slet"],
            ["Helligdage",        "holidays",               "Auto-generer for år, opret manuelt, slet. Kræver 'manage_holidays'-rettighed."],
            ["Disponentgrupper",  "dispatcher_groups",      "Opret, omdøb/rediger beskrivelse, slet (fjerner automatisk tilknytning hos medlemmer)"],
            ["Aftale",            "master_agreement_kinds", "Opret nye, rediger label/aktiv/kræver overenskomsttype. De to systemtyper (hourly_fixed/hourly_flexible) kan ikke slettes; nye typer kan slettes hvis ingen medarbejder bruger dem."],
        ]
    )
    note_box(doc,
        "Alle handlinger i Stamdata-modulet kræver 'stamdata'-rettighed. "
        "Helligdage-fanen kræver desuden 'manage_holidays'-rettighed.",
        "BEMÆRK"
    )

    heading(doc, "Helligdage – API-endepunkter", 2, "6.3")
    header_table(doc,
        ["Endepunkt", "Metode", "Rettighed", "Beskrivelse"],
        [
            ["GET /api/stamdata/holidays",                  "GET",    "stamdata",        "Henter alle helligdage. Valgfri ?year=YYYY-filter."],
            ["POST /api/stamdata/holidays",                 "POST",   "manage_holidays", "Opretter helligdag. Body: date (YYYY-MM-DD), name, half_day_from (opt.)."],
            ["DELETE /api/stamdata/holidays/{id}",          "DELETE", "manage_holidays", "Sletter helligdag med angivet id."],
            ["POST /api/stamdata/holidays/generate/{year}", "POST",   "manage_holidays", "Auto-genererer helligdage for ét år (idempotent). År: 2020–2100."],
        ]
    )
    body(doc, (
        "Helligdagsdata bruges i aktivitetskalenderens frontend: "
        "loadHolidaysForPeriod(startDate, endDate) kaldes parallelt med loadActivities() og "
        "gemmer resultatet i state.holidays[]. renderActivitiesTable() markerer helligdagskolonner "
        "med baggrundsfarven #056a10 (mørkegrøn). Halvdagshelligdage vises med et '½ fra HH:MM'-badge "
        "og helligdagsnavnet som tooltip. Hvis dagens dato falder på en helligdag, bevares "
        "today-indikatoren via border-bottom i accentfarven."
    ))

    heading(doc, "Medarbejdertillæg (employee_supplements)", 2, "6.4")
    body(doc, (
        "Ud over overenskomsttypens timesats kan hver medarbejder have et individuelt kr/time-tillæg "
        "('Ikke overenskomstmæssigt tillæg'), historisk sporet i en selvstændig tabel employee_supplements. "
        "Tabellen administreres fra det selvstændige sidebar-punktet 'Tillæg' (kræver rettigheden "
        "manage_employee_supplements) og påvirker selve lønberegningen – ikke kun visning."
    ))
    header_table(doc,
        ["Felt", "Type", "Beskrivelse"],
        [
            ["employee_id",  "FK → Employee",  "Tilknyttet medarbejder"],
            ["name",         "String",         "Altid 'Ikke overenskomstmæssigt tillæg' – hardkodet server-side, ikke redigerbar"],
            ["type",         "String",         "Altid 'Timebaseret' – hardkodet server-side"],
            ["value",        "Numeric(10,2)",  "Tillæggets værdi i kr/time. Skal være > 0, afrundes til 2 decimaler før gemning"],
            ["start_date",   "Date",           "Gyldighedsperiodens start – default dags dato ved oprettelse"],
            ["end_date",     "Date",           "Gyldighedsperiodens slut – default 9999-12-31 (åbentstående)"],
        ]
    )
    body(doc, (
        "Status (Aktiv/Inaktiv) er IKKE et lagret felt – det beregnes ved hver visning ud fra om "
        "dags dato ligger i intervallet [start_date, end_date]. Et tillæg med fremtidig startdato "
        "vises derfor som Inaktiv, indtil startdatoen er nået."
    ))

    heading(doc, "Livscyklus og satsopslag", 2, "6.4.1")
    for step in [
        "Oprettelse (POST /api/employee-supplements): den medarbejders eksisterende åbentstående række (end_date = 9999-12-31) får automatisk sin end_date sat til ny_start_dato − 1 dag, og en ny åbentstående række indsættes. Ny start_date skal være efter den forrige rækkes start_date, ellers afvises oprettelsen.",
        "Afslutning (POST /api/employee-supplements/{id}/end): sætter IKKE end_date til dags dato, men til slutdatoen for den lønperiode dags dato falder i (get_or_create_period_for_date()). Tillægget gælder derfor stadig resten af den igangværende lønperiode og bortfalder først fra den efterfølgende periode.",
        "Der findes ingen redigerings- eller sletningsmulighed for eksisterende rækker – kun oprettelse af nye og afslutning af den aktive.",
        "Et partielt unikt indeks (uq_employee_supplements_one_open_row) sikrer på databaseniveau at en medarbejder kun kan have ÉN åbentstående række ad gangen (WHERE end_date = '9999-12-31') – forhindrer et race condition ved samtidige oprettelser.",
    ]:
        bullet(doc, step)
    body(doc, (
        "Ved lønberegning slår get_active_supplement_for_period(db, employee_id, periode_start, periode_slut) "
        "(calculators/rates_loader.py) op efter den række hvis gyldighedsperiode OVERLAPPER den beregnede periode. "
        "Overlapper flere rækker (fordi et nyt tillæg er oprettet midt i en periode), vinder rækken med nyeste "
        "start_date – for hele perioden (ingen dag-for-dag splitning). Denne regel gør samtidig en genberegning "
        "af en gammel, afsluttet periode historisk korrekt: det tillæg der dengang overlappede perioden, ændres "
        "ikke bagudrettet af senere oprettede tillæg."
    ))
    note_box(doc,
        "hourly_rate forhøjes med tillæggets value umiddelbart efter det almindelige overenskomstopslag i "
        "_calculate_employee() (payroll_router.py) – ÉN variabel der allerede bruges alle steder nedstrøms "
        "(normaltid/kode 1, SH-betaling, CSV-rækkerne for AFSPADSERING/SYGDOM/BARSEL/SKOLE_KURSUS m.fl.). "
        "Samme funktion genbruges af absence_overview_router.py, så Fraværsoversigtens viste sats altid stemmer "
        "overens med den sats der reelt udbetales i lønkørslen.",
        "TEKNISK NOTE"
    )

    heading(doc, "API-endepunkter", 2, "6.4.2")
    header_table(doc,
        ["Endepunkt", "Metode", "Beskrivelse"],
        [
            ["/api/employee-supplements",              "GET",  "Liste, filtreret på employee_id/from/to (gyldighedsperiode-overlap)."],
            ["/api/employee-supplements/active/{id}",   "GET",  "Det aktuelt aktive tillæg for én medarbejder, eller null."],
            ["/api/employee-supplements",               "POST", "Opret nyt tillæg (se livscyklus ovenfor)."],
            ["/api/employee-supplements/{id}/end",      "POST", "Afslut det aktuelt aktive tillæg fra og med den efterfølgende lønperiode."],
        ]
    )
    body(doc, (
        "Alle fire endepunkter kræver rettigheden manage_employee_supplements. Medarbejder-modalen "
        "(fanen 'Medarbejdere') viser tillæggets aktuelle værdi i et read-only felt under Overenskomsttype "
        "(hentet via GET /active/{id}) – feltet kan ikke redigeres eller udfyldes derfra, kun under fanen 'Tillæg'."
    ))

    # ── 7. Lønkørsel ─────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Lønkørsel", 1, "7")
    body(doc, (
        "Lønkørslen foregår i routers/payroll_router.py. KUN aktiviteter med status "
        "'approved' indgår i beregningen."
    ))

    heading(doc, "API-endepunkter", 2, "7.1")
    header_table(doc,
        ["Endepunkt", "Metode", "Beskrivelse"],
        [
            ["/api/payroll/preview",         "POST", "Returnerer JSON med mellemregninger for alle eller én medarbejder."],
            ["/api/payroll/proevekoersel",   "POST", "Genererer Excel-fil med prøvekørsel og returnerer download-link."],
            ["/api/payroll/proevekoersel-gem","POST","Gemmer Excel-fil i brugervalgt mappe (tkinter-dialog)."],
            ["/api/payroll/export-csv",      "POST", "Genererer Danløn CSV-fil og downloader den. Afviser (400) hvis perioden allerede er låst, eller der er afventende aktiviteter i perioden."],
            ["/api/payroll/pdf-timesedler",  "POST", "Genererer PDF-timesedler og gemmer i valgt mappe."],
            ["/api/payroll/browse-folder",   "GET",  "Åbner Windows-mappe-dialog (tkinter) og returnerer valgt sti."],
            ["/api/payroll/downloads-folder","GET",  "Returnerer stien til brugerens Downloads-mappe."],
        ]
    )

    heading(doc, "Beregningslogik (_calculate_employee)", 2, "7.2")
    for step in [
        "Itererer alle 14 dage i perioden – ikke kun dage med aktiviteter.",
        "Dage uden aktivitet returneres med 0 timer på alle kolonner.",
        "Overnatning (activity_type = 'overnatning'): tæller forekomster som overnight_count og springer dagslinje-logikken over (continue). Eksponeres i resultatet som overnight_count, overnight_rate og overnight_kr (flat sats pr. forekomst fra Salttillæg og overnatning.xlsx).",
        "Fraværsdage vises med typenavn i stedet for timer i dagsoversigten, men timerne beregnes og summeres stadig separat pr. fraværstype (fx sygdom_hours, afspadsering_hours, feriefri_hours, barsel_hours) og indgår i Danløn CSV'en, hvis løntypen er sat til 'Medtag i CSV' i Stamdata. Ferie er en undtagelse: den tælles, men er som default ekskluderet fra CSV'en.",
        "For normale aktiviteter: kalder calculate_overtime() med aktivitetens start/slut, pauseintervaller og normaltimer for den pågældende dag.",
        "Summerer normal tid, OT_BEFORE, OT_13 og OT_EXTRA timer samt pålæsning/aflæsning.",
        "Beregner kr. = timer × timesats (normal) + tillægstimer × satser.",
        "Returnerer dict med totaler og linjeposter pr. dag (alle 14 dage i perioden).",
    ]:
        bullet(doc, step)

    heading(doc, "Excel-formatering i prøvekørsel", 2, "7.3")
    body(doc, "Excel-filen der genereres ved prøvekørsel har følgende formatering:")
    two_col_table(doc, [
        ["Grøn første række",    "Første datalinje for hver medarbejder markeres med grøn baggrund (farve #C6EFCE) for nem identifikation af ny medarbejder."],
        ["Tom skillerække",      "Efter hver medarbejders totallinje indsættes en tom række som visuel adskillelse."],
        ["Frossen headerrække",  "Headerrækken (række 1) er frossen og følger med ved lodret scroll."],
        ["Fraværstyper",         "Fraværsdage (ferie, fri, afspadsering, skole/kursus) vises med typenavn i Normal tid-kolonnen; øvrige kolonner er blanke."],
        ["Nul-dage",             "Dage uden aktivitet vises med 0 på alle talkolonner."],
    ])

    heading(doc, "Danløn CSV-format", 2, "7.4")
    body(doc, (
        "Kolonnerne i den eksporterede CSV-fil (semikolon-separeret). "
        "Kolonne 5 og 6 er valgfrie per løntypekode og styres via Stamdata:"
    ))
    header_table(doc,
        ["Kolonne", "Altid med?", "Indhold"],
        [
            ["CVR",           "Ja",      "Virksomhedens CVR-nummer"],
            ["Medarbejdernr", "Ja",      "Lønnummer fra medarbejderkortet"],
            ["Lønkode",       "Ja",      "Danløn-kode for lønarten (konfigureres i Stamdata)"],
            ["Antal",         "Ja",      "Timer (2 decimaler) eller antal forekomster (heltal) afhængigt af Antal-type"],
            ["Sats",          "Valgfri", "Sats i kr. (2 decimaler). Slås fra ved 'Inkluder sats' = Nej i Stamdata"],
            ["Total",         "Valgfri", "Antal × Sats (2 decimaler). Tilføjes ved 'Inkluder total' = Ja i Stamdata"],
        ]
    )
    body(doc, (
        "CSV-kolonneopsætningen konfigureres per løntypekode i Stamdata → Løntypekoder: "
        "Antal-type (Timer/Antal), Sats-kilde (timesats, overtidssatser, salttillæg, overnatning, dagpenge), "
        "Inkluder sats og Inkluder total."
    ))

    # ── 8. Aktivitetshåndtering ───────────────────────────────────────────
    heading(doc, "Aktivitetshåndtering", 1, "8")

    heading(doc, "Statusflow", 2, "8.1")
    body(doc,
        "pending → approved (godkendt – approved_by sættes til brugerens initialer) → "
        "deactivated (deaktiveret – deactivated_by sættes til brugerens initialer).\n"
        "Fraværstyper (ferie, sygdom, afspadsering, skole/kursus mv.) oprettes direkte som approved, med approved_by sat til den opretstende brugers initialer (ikke pending, da de ikke skal manuelt godkendes bagefter).\n"
        "approved_by og deactivated_by er separate felter – approved_by bruges udelukkende til godkendelser."
    )

    heading(doc, "Rettelse og fortryd", 2, "8.2")
    for step in [
        "Første rettelse: original_start_time og original_end_time gemmes automatisk.",
        "Efterfølgende rettelser: originaltiderne overskrives ikke.",
        "Fortryd: nulstiller start_time og end_time til de gemte originaler.",
    ]:
        bullet(doc, step)

    heading(doc, "Opdeling (split)", 2, "8.3")
    body(doc, (
        "En aktivitet kan opdeles ved et valgt tidspunkt. Systemet opretter to "
        "børneaktiviteter med parent_activity_id og split_part (1 og 2). "
        "Pauser og segmenter fordeles proportionalt til den korrekte del."
    ))

    heading(doc, "Aktivitetsflags", 2, "8.4")
    two_col_table(doc, [
        ["is_under_4h",  "Aktiviteten er under 4 timer – kommentar påkrævet ved godkendelse."],
        ["is_over_12h",  "Aktiviteten er over 12 timer – vises som advarsel i UI."],
        ["is_manual",    "Aktiviteten er manuelt oprettet (source=manual, ikke fra tachograf)."],
        ["is_edited",    "Tiderne er blevet ændret fra tachografdata."],
        ["has_split_children", "Aktiviteten er opdelt i to børneaktiviteter."],
    ])

    # ── 9. Frontend ───────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Frontendarkitektur", 1, "9")
    body(doc, (
        "Hele frontenden er én HTML-fil (templates/index.html) med én JavaScript-fil "
        "(static/js/app.js). Der bruges ingen React, Vue eller andre frameworks."
    ))

    heading(doc, "State-management", 2, "9.1")
    body(doc, "Et globalt state-objekt holder applikationens tilstand:")
    two_col_table(doc, [
        ["state.employees",      "Liste af alle aktive medarbejdere (inkl. dispatcher_group, work_schedule osv.)"],
        ["state.activities",     "Alle aktiviteter for den valgte periode"],
        ["state.periodInfo",     "Aktuel periodeinfo: start/slut, tæller for pending/approved/deactivated"],
        ["state.agreementTypes", "Overenskomsttyper og satser"],
        ["state.splitMin/Max",   "Grænser for opdeling (gemmes i state for validering)"],
        ["manualPauses",         "Pauseintervaller for aktivitet under oprettelse: liste af [ISO-start, ISO-slut]-par. Nulstilles ved åbning af oprettelsesmodalen."],
    ])

    heading(doc, "Cache-busting", 2, "9.2")
    body(doc, (
        "main.py beregner app.js's ændringstidspunkt (st_mtime) og indsætter det som "
        "forespørgselsparameter: <script src='/static/js/app.js?v={mtime}'>. "
        "Browsercachen brydes automatisk ved hver ændring af filen."
    ))

    heading(doc, "Datetime-picker", 2, "9.3")
    body(doc, (
        "En brugerdefineret datetime-komponent erstatter browserens native datetime-local "
        "scroll-input. Komponenten består af: date-input (dato) + number-input (time, 0-23) "
        "+ ':' + number-input (minut, 0-59). Alle styles er inline i JavaScript for at "
        "undgå browsercache-problemer."
    ))
    two_col_table(doc, [
        ["buildDatetimePicker(id, isoValue)", "Bygger komponenten og indsætter den i elementet med det givne id."],
        ["readDatetimePicker(id)",            "Læser de tre felter og returnerer ISO-datetime-streng."],
        ["setDatetimePicker(id, isoValue)",   "Sætter alle tre felter fra en ISO-datetime-streng."],
    ])

    heading(doc, "Date-picker (datovalg)", 2, "9.4")
    body(doc, (
        "En brugerdefineret dato-komponent (buildDatePicker) erstatter browserens native "
        "date-input på steder hvor det er vigtigt at kunne vælge årstal fra dropdown. "
        "Komponenten består af et tekstfelt der åbner et popup-kalender med måned-dropdown, "
        "årstal-dropdown (1950–2100) og en klassisk månedsgrid med klikbare datoer."
    ))
    two_col_table(doc, [
        ["buildDatePicker(id, initialValue)", "Bygger komponenten i elementet med det givne id. Accepts ISO-dato eller tom streng."],
        ["readDatePicker(id)",                "Returnerer den valgte dato som ISO-streng (YYYY-MM-DD)."],
        ["setDatePicker(id, iso)",            "Opdaterer display og hidden-felt til en ny ISO-dato. Opdaterer også kalender-griddet."],
    ])
    body(doc, "Bruges på to steder i applikationen:")
    bullet(doc, "Medarbejdermodal: 'Ansættelsesdato' og 'Fratrædelsesdato'.")
    bullet(doc, "Aktivitetsoversigt: Datovælgeren i periodelinjen (hop til dato).")
    note_box(doc,
        "Popup'en bruger position:fixed og getBoundingClientRect() for korrekt placering "
        "inde i scrollbare modaler. Specialtilfældet 9999-12-31 (ingen fratrædelse) "
        "vises som '31-12-9999' og åbner kalenderen ved indeværende år.",
        "TEKNISK NOTE"
    )

    heading(doc, "Filtrering", 2, "9.5")
    body(doc, "Aktivitetstabellen har tre filtre der kombineres:")
    two_col_table(doc, [
        ["Status-filter",      "Alle / Afventer / Godkendt / Deaktiveret"],
        ["Medarbejder-filter", "Viser kun aktiviteter for valgte medarbejder. Opdateres automatisk når afdelingsfilter ændres."],
        ["Afdelings-filter",   "Filtrerer aktivitetstabellen OG medarbejder-dropdown til kun at vise medarbejdere i den valgte disponentgruppe."],
    ])

    heading(doc, "Stamdata-view", 2, "9.6")
    body(doc, (
        "Stamdata-view aktiveres fra menupunktet '⚙️ Stamdata' i venstre menu (kræver 'stamdata'-rettighed). "
        "View'et indeholder en tab-navigator med ni faner – kun én pane er synlig ad gangen:"
    ))
    two_col_table(doc, [
        ["switchStamdataTab(tab)", "Skifter aktiv pane og opdaterer fane-styling (border, farve, vægt)."],
        ["loadStamdata()",         "Kaldes fra setView('stamdata') og indlæser data til alle ni faner parallelt."],
        ["btn-stamdata-add-*",     "'+Tilføj'-knap i toolbar vises kun for den aktive fane."],
    ])
    body(doc, (
        "Fraværstyper-fanen viser alle rækker fra master_absence_types med Aktiv-badge (grøn/grå). "
        "Alle typer har Rediger- og Slet-knap. "
        "Løntypekoder-fanen giver mulighed for at redigere type (label), Danløn-kode, Medtag i CSV, "
        "Antal-type (Timer/Antal), Sats-kilde, Inkluder sats og Inkluder total samt slette alle koder. "
        "Alle disse felter kan også sættes ved oprettelse af en ny løntypekode."
    ))
    body(doc, (
        "Aftale-fanen viser rækkerne fra master_agreement_kinds med badges for Type "
        "(System/Brugeroprettet), Kræver overenskomsttype (Ja/Nej) og Aktiv. Systemtyperne har "
        "kun en Rediger-knap (label/aktiv/kræver overenskomsttype); brugeroprettede typer har "
        "desuden en Slet-knap. loadStamdataAgreementKinds() genindlæser efter hver ændring også "
        "state.agreementKinds via GET /api/employees/agreement-kinds, så medarbejder-modalens "
        "'Aftale'-dropdown (fillAgreementKindSelect()) altid er opdateret uden sideopdatering. "
        "onAgreementKindChange() slår den valgte types requires_agreement_type op og viser/skjuler "
        "den røde stjerne ved Overenskomsttype-feltet i medarbejder-modalen."
    ))

    heading(doc, "Pausehåndtering i oprettelsesmodalen", 2, "9.7")
    body(doc, (
        "Når en manuel aktivitet oprettes, kan der tilføjes et vilkårligt antal pauser via "
        "'+ Tilføj pause'-knappen. Pauserne gemmes i databasefeltet pause_intervals og fratrækkes "
        "automatisk i lønberegning, varighedsvisning og aktivitetsfordeling."
    ))
    header_table(doc,
        ["Funktion (app.js)", "Beskrivelse"],
        [
            ["addManualPause()",      "Læser startdato fra aktivitetens starttid og åbner modal-pause med dt-picker (datofeltet skjult – kun HH:MM vises)."],
            ["confirmPause()",        "Validerer at slut > start, tilføjer [ISO-start, ISO-slut] til manualPauses[] og lukker modal-pause."],
            ["deleteManualPause(idx)","Fjerner pause ved index fra manualPauses[] og opdaterer listen."],
            ["renderManualPauses()", "Tegner #manual-pauses-list med pause-badges og ×-slet-knapper."],
        ]
    )
    body(doc, "Pauseinterval-flow ved POST af aktivitet:")
    for step in [
        "manualPauses[] sendes som pause_intervals i POST-body til POST /api/activities.",
        "ActivityCreate-skemaet validerer listen (list[list[str]]).",
        "_duration_minutes(a) i activities.py klipper hvert interval til aktivitetens bounds og fratrækker det fra bruttotiden. Påvirker is_under_4h og is_over_12h.",
        "calculate_overtime() i overtime.py bruger _subtract_pauses() til at fratrække pauserne i det korrekte tillægsvindue.",
        "openActivityDetail() i app.js beregner Aktivitetsfordeling-bjælken: for manuelle aktiviteter uden tachografsegmentdata beregnes Hvil/pause-% fra pause_intervals og Effektiv tid-% som resten.",
    ]:
        bullet(doc, step)
    note_box(doc,
        "Pausemodalen (modal-pause) bruger buildDatetimePicker med datofeltet skjult – "
        "brugeren ser kun HH:MM. Datoen arves fra aktivitetens startdato.",
        "TEKNISK NOTE"
    )

    # ── 10. Drift og vedligehold (FAQ) ────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Drift og vedligehold – ofte stillede spørgsmål", 1, "10")

    heading(doc, "Hvor er databasen, og hvordan kører den?", 2, "10.1")
    body(doc, (
        "Databasen er én enkelt fil: app/database/lonsystem.db. "
        "SQLite er ikke et separat program der skal startes – Python-koden "
        "læser og skriver direkte i filen. Ingen databaseserver, ingen porte, "
        "ingen baggrundstjeneste. Systemet opretter automatisk en ny, tom database "
        "første gang det startes, hvis filen ikke eksisterer."
    ))
    two_col_table(doc, [
        ["Placering",   "app/database/lonsystem.db"],
        ["Størrelse",   "Starter ved ~140 KB, vokser efterhånden som data opbygges"],
        ["Backup",      "Kopiér filen (mens serveren er stoppet) for at tage backup"],
        ["Nulstilling", "Slet filen – systemet opretter en tom database ved næste opstart"],
    ])

    note_box(doc,
        "Filen bør IKKE ligge i en OneDrive- eller Dropbox-mappe når systemet er i produktion. "
        "Cloud-synkronisering og SQLite konflikter og kan i sjældne tilfælde korruptere databasen. "
        "På en produktionsserver skal filen ligge lokalt på serveren.",
        "VIGTIGT"
    )

    heading(doc, "Hvordan starter og stopper jeg serveren?", 2, "10.2")
    body(doc, "Serveren (uvicorn) er den proces der holder systemet kørende og svarer på browser-forespørgsler.")

    p = doc.add_paragraph()
    r = p.add_run("Stop serveren")
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11)
    body(doc, "Tryk Ctrl+C i det terminalvindue hvor uvicorn kører.", space_after=2)
    body(doc, "Alternativt: Åbn Task Manager → find python.exe → Afslut opgave.", space_after=6)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Start serveren")
    r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(11)

    body(doc, "Der er to måder at starte serveren på, afhængigt af situationen:", space_after=4)

    body(doc, "Daglig drift (anbefalet):", space_after=2)
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(1)
    code.paragraph_format.space_after = Pt(4)
    cr = code.add_run(
        'python.exe -m uvicorn --app-dir app main:app --host 0.0.0.0 --port 8000'
    )
    cr.font.name = "Courier New"; cr.font.size = Pt(9)
    body(doc, "Hurtig og stabil. Ændringer i Python-filer kræver manuel servergenstart (Ctrl+C → start igen).", space_after=8)

    body(doc, "Under udvikling (når programmet ændres af en udvikler):", space_after=2)
    code2 = doc.add_paragraph()
    code2.paragraph_format.left_indent = Cm(1)
    code2.paragraph_format.space_after = Pt(4)
    cr2 = code2.add_run(
        'python.exe -m uvicorn --app-dir app main:app --host 0.0.0.0 --port 8000 --reload'
    )
    cr2.font.name = "Courier New"; cr2.font.size = Pt(9)
    body(doc, (
        "Med --reload genstarter serveren automatisk hver gang en Python-fil gemmes, "
        "så ændringer træder i kraft ved næste browser-reload uden manuel genstart. "
        "Lidt langsommere end daglig drift, men bekvem under aktiv udvikling."
    ), space_after=8)

    body(doc, (
        "Når serveren kører, er systemet tilgængeligt i browseren på "
        "http://127.0.0.1:8000 (lokal maskine) eller serverens IP-adresse "
        "på netværket."
    ))

    note_box(doc,
        "Ændringer i JavaScript-, CSS- og HTML-filer (app.js, style.css, index.html) "
        "træder altid i kraft ved browser-reload – uanset om serveren kører med eller uden --reload. "
        "Det er kun ændringer i Python-filer (.py) der kræver servergenstart.",
        "BEMÆRK"
    )

    heading(doc, "Hvad kræves for at flytte systemet til en fælles server?", 2, "10.3")
    body(doc, (
        "Da systemet er bygget som en webapplikation er det designet til delt brug. "
        "Alle ansatte tilgår det via en browser – ingen installation på klientmaskinen. "
        "Flytningen til en fælles server kræver:"
    ))
    for step in [
        "En dedikeret maskine (eksisterende server, spare-PC eller cloud-VM) der er tændt i arbejdstiden.",
        "Python 3.13 installeret på serveren.",
        "Koden og databasefilen kopieret til serveren (uden for en cloud-synkroniseret mappe).",
        "En fast lokal IP-adresse på servermaskinen (konfigureres i router eller af IT).",
        "Serveren sat op til at starte uvicorn automatisk ved opstart – anbefalet via NSSM (gratis Windows-værktøj der kører uvicorn som en Windows-tjeneste).",
        "Ansatte åbner browseren og går til serverens IP-adresse og port, f.eks. http://192.168.1.50:8000.",
    ]:
        bullet(doc, step)

    note_box(doc,
        "Hvis systemet skal tilgås uden for virksomhedens netværk (hjemmefra, fra mobil), "
        "kræves enten VPN-adgang til netværket eller en cloudserver med HTTPS-certifikat. "
        "Dette er en IT-beslutning der også indebærer en GDPR-vurdering af datalokation.",
        "BEMÆRK"
    )

    heading(doc, "Hvordan ryddes testdata og systemet gøres klar til produktion?", 2, "10.4")
    body(doc, "Følgende fremgangsmåde sikrer et rent produktionssystem:")
    for i, step in enumerate([
        "Stop serveren (Ctrl+C eller Task Manager).",
        "Notér alle rigtige medarbejdere der skal oprettes: lønnumre, førerkortnumre, overenskomsttyper, timefordeling og afdelinger.",
        "Slet filen app/database/lonsystem.db.",
        "Kontrollér at Excel-filerne (Overtid satser.xlsx, Overenskomsttyper og timesatser.xlsx, Fraværstyper.xlsx) indeholder de korrekte satser – de bruges som udgangspunkt for seeding af Stamdata-tabellerne.",
        "Start serveren igen – systemet opretter automatisk en tom database og seeder Stamdata-tabellerne fra Excel-filerne.",
        "Kontrollér og juster om nødvendigt satser og fraværstyper i Stamdata-modulet (⚙️ Stamdata i menuen).",
        "Opret de rigtige medarbejdere i systemet.",
        "Systemet er klar til produktionsbrug.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")

    body(doc, (
        "Excel-filerne bruges kun til den initiale seeding. Efterfølgende ændringer i satser og "
        "fraværstyper skal foretages via Stamdata-modulet i brugerfladen – ikke i Excel-filerne."
    ))

    # ── 11. Backup-utility ────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Backup-utility", 1, "11")
    body(doc, (
        "Systemet leveres med et separat backup-utility der kører automatisk 4 gange "
        "dagligt og bevarer de seneste 5 dages backup-historik. Utility'et ligger i "
        "backup/-mappen ved siden af app/-mappen."
    ))

    heading(doc, "Filer", 2, "11.1")
    two_col_table(doc, [
        ["backup/backup.py",       "Selve backup-scriptet. Kan køres manuelt eller via Task Scheduler."],
        ["backup/installer.ps1",   "PowerShell-script der registrerer backup-opgaven i Windows Task Scheduler."],
        ["backup/arkiv/",          "Mappe med ZIP-filer – én pr. kørsel."],
        ["backup/backup.log",      "Logfil med tidsstempler og resultater for alle kørsler."],
    ])

    heading(doc, "Hvad der tages backup af", 2, "11.2")
    two_col_table(doc, [
        ["app/database/lonsystem.db",                    "Databasen – det vigtigste. Kopieres via SQLites eget backup-API, sikkert mens systemet kører."],
        ["app/Overtid satser.xlsx",                      "Overtidssatser."],
        ["app/Overenskomsttyper og timesatser.xlsx",     "Timesatser og overenskomsttyper."],
        ["app/Afdelinger.xlsx",                          "Historisk Excel-fil – disponentgrupper læses ikke længere herfra, men fra databasetabellen dispatcher_groups (se afsnit 2.1)."],
    ])
    body(doc, (
        "Hvert backup gemmes som en ZIP-fil med navneformatet lonsystem_YYYY-MM-DD_HH-MM.zip. "
        "ZIP-filer ældre end 5 dage slettes automatisk. Med 4 kørsler dagligt i 5 dage "
        "er der til enhver tid 20 backups i arkivet."
    ))

    heading(doc, "Teknisk detalje – sikker kopiering", 2, "11.3")
    body(doc, (
        "SQLite-databasen kopieres via Pythons sqlite3.backup()-metode. Denne metode "
        "bruger SQLites interne backup-API og producerer en konsistent snapshot selv "
        "om der skrives til databasen samtidig (WAL-mode). Det er ikke sikkert blot "
        "at kopiere .db-filen direkte med shutil.copy() mens systemet kører."
    ))

    heading(doc, "Installation", 2, "11.4")
    body(doc, "Kør installer.ps1 én gang som Administrator:")
    for i, step in enumerate([
        "Søg på 'PowerShell' i Start-menuen.",
        "Højreklik → Kør som administrator.",
        'Kør: & "...\\backup\\installer.ps1"',
        "Test med: Start-ScheduledTask -TaskName \"Lønsystem Backup\"",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")

    body(doc, (
        "Opgaven kører som NT AUTHORITY\\SYSTEM og kræver ingen gemt adgangskode. "
        "Den kører selv når ingen bruger er logget ind på serveren."
    ))

    heading(doc, "Gendannelse fra backup", 2, "11.5")
    for i, step in enumerate([
        "Stop serveren (uvicorn).",
        "Pak den ønskede ZIP-fil ud.",
        "Kopier lonsystem.db til app/database/lonsystem.db (erstat den eksisterende).",
        "Kopier eventuelt Excel-filerne tilbage til app/ hvis de også er beskadigede.",
        "Start serveren igen.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")

    # ── 12. Netværksadgang og adgang fra andre maskiner ────────────────────
    heading(doc, "Netværksadgang fra andre maskiner", 1, "12")
    body(doc, (
        "Systemet er designet til at køre på én server og tilgås fra mange klientmaskiner "
        "via en webbrowser. Der kræves ingen installation på klientmaskinerne."
    ))

    heading(doc, "Forudsætninger", 2, "12.1")
    two_col_table(doc, [
        ["--host 0.0.0.0",          "Uvicorn skal lytte på alle netværksgrænseflader (ikke kun 127.0.0.1). Sat i .claude/launch.json."],
        ["Windows Firewall",        "Port 8000 skal åbnes for indgående TCP-trafik på det private netværk."],
        ["Fast lokal IP-adresse",   "Servermaskinens IP bør være fast (konfigureres i router eller netværksindstillinger)."],
    ])

    heading(doc, "Åbn port 8000 i Windows Firewall", 2, "12.2")
    body(doc, "Kør én gang som Administrator i PowerShell:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(1)
    code.paragraph_format.space_after = Pt(8)
    cr = code.add_run(
        'New-NetFirewallRule -DisplayName "Lønsystem port 8000" '
        '-Direction Inbound -Protocol TCP -LocalPort 8000 '
        '-Action Allow -Profile Private'
    )
    cr.font.name = "Courier New"; cr.font.size = Pt(9)

    heading(doc, "Find serverens IP-adresse", 2, "12.3")
    body(doc, "Kør ipconfig i PowerShell og aflæs IPv4 Address under det aktive netværkskort:")
    code2 = doc.add_paragraph()
    code2.paragraph_format.left_indent = Cm(1)
    code2.paragraph_format.space_after = Pt(8)
    cr2 = code2.add_run("ipconfig | Select-String 'IPv4'")
    cr2.font.name = "Courier New"; cr2.font.size = Pt(9)

    heading(doc, "Tilgang fra klientmaskine", 2, "12.4")
    body(doc, (
        "Åbn en webbrowser på klientmaskinen og gå til serverens IP-adresse og port 8000. "
        "Eksempel hvis serverens IP er 192.168.1.29:"
    ))
    code3 = doc.add_paragraph()
    code3.paragraph_format.left_indent = Cm(1)
    code3.paragraph_format.space_after = Pt(8)
    cr3 = code3.add_run("http://192.168.1.29:8000")
    cr3.font.name = "Courier New"; cr3.font.size = Pt(9)

    note_box(doc,
        "Adressen 127.0.0.1 eller localhost virker kun på selve servermaskinen. "
        "Andre maskiner skal bruge serverens rigtige IP-adresse på netværket.",
        "OBS"
    )

    # ── 13. Lønafregning ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Lønafregning", 1, "13")
    body(doc, (
        "Lønafregning (routers/payroll_settlement_router.py) er en selvstændig fane placeret "
        "under 'Lønkørsel' i venstre menu. Den viser periodetotaler og en 14-dages tabel pr. "
        "medarbejder, og genbruger _calculate_employee() (payroll_router.py) som eneste "
        "beregningskilde – ingen lønmatematik er duplikeret."
    ))
    body(doc, (
        "I modsætning til de fleste andre visninger har Lønafregning ingen egen "
        "periode-navigation. Siden følger i stedet den globale periode (state.currentPeriodStart), "
        "som vælges via frem/tilbage-pilene i Aktivitetsoversigten – nøjagtig samme mekanisme "
        "som Lønkørsel allerede bruger. Uden valgt periode falder den tilbage til dagens periode."
    ))

    heading(doc, "API-endepunkter", 2, "13.1")
    header_table(doc,
        ["Endepunkt", "Metode", "Beskrivelse"],
        [
            ["/api/payroll-settlement/preview",          "GET",  "JSON med periodetotaler og pr.-medarbejder headline + 14-dages tabel. Parameter period_start (valgfri, default dagens periode)."],
            ["/api/payroll-settlement/downloads-folder", "GET",  "Returnerer stien til brugerens Downloads-mappe som forslag."],
            ["/api/payroll-settlement/browse-folder",    "GET",  "Åbner Windows-mappe-dialog (tkinter) og returnerer valgt sti."],
            ["/api/payroll-settlement/export-csv",       "POST", "Genererer og gemmer CSV-filen i valgt mappe. Kræver låst periode – se afsnit 13.3."],
        ]
    )
    body(doc, (
        "Alle fire endepunkter kræver rettigheden payroll_settlement_view; export-csv kræver "
        "derudover payroll_settlement_export. Se afsnit 11.3/rettighedsoversigten i Brugervejledningen."
    ))

    heading(doc, "Beregningslogik og fraværsbetaling", 2, "13.2")
    body(doc, (
        "_employee_settlement_data() (payroll_settlement_router.py) kalder _calculate_employee() "
        "og lægger to ting oveni: (1) overenskomstsats og personligt tillæg vises som to adskilte "
        "tal i stedet for kun den kombinerede hourly_rate, og (2) springertillæg (normalt usynligt "
        "i Lønkørsel-visningen) lægges til medarbejderens totale løn med eget kr-beløb."
    ))
    body(doc, (
        "Fraværstyper med en etableret betalingsregel viser deres egne timer og beløb i dagens "
        "række (i stedet for 0,00) OG tæller med i medarbejderens 'Total løn' samt periodens "
        "samlede total. Følgende typer er dækket, hver med sin egen linje i topsummeringen "
        "'Total sum for perioden' (kun vist hvis beløbet er større end 0 kr):"
    ))
    header_table(doc,
        ["Fraværstype", "Sats der bruges"],
        [
            ["Sygdom, Barn 1.sygedag, Graviditetsbetinget sygdom", "Medarbejderens timesats (hourly_rate)"],
            ["§56 syg, Barn 1.sygedag u. 8 uger", "Dagpengesats (Stamdata → Tillæg, 'Dagpenge §56')"],
            ["Sygdom u. 8 uger", "Ulønnet – altid 0 kr., men timerne vises"],
            ["Barsel, Feriefri, Ferie, Skole/kursus, Afspadsering", "Medarbejderens timesats (hourly_rate)"],
        ]
    )
    note_box(doc,
        "Feriefri vises som ÉN samlet linje uanset om medarbejderen er fuldlønnet eller "
        "timelønnet – emp.fuldloennet afgør i den almindelige Danløn-CSV kun hvilken kode "
        "(FERIEFRI_FULDLOENNET/-TIMELOENNET) timerne rapporteres under, ikke selve beløbet, "
        "som beregnes ens i begge tilfælde (hours × hourly_rate). Bekræftet af bruger 2026-08-25.",
        "TEKNISK NOTE"
    )
    body(doc, (
        "Selvbetalt fridag, Barn 2-3.sygedag, Barsel u. løn og typen 'Fri' har ingen etableret "
        "betalingsregel i systemet (samme som i den almindelige Danløn CSV) og har derfor hverken "
        "beløb i dagsrækken eller egen linje i topsummeringen."
    ))
    body(doc, (
        "'Total uden fravær' er en delsum der lægger Grundtimeløn inkl. tillæg, Overtid Timen "
        "før, Overtid 1-3 time efter og Øvrig overtid sammen – IKKE salttillæg eller nogen "
        "fraværstype. Den samlede 'Total sum for denne periode' er summen af alle medarbejderes "
        "fulde total_kr (arbejdstid + salt + springer + al fraværsbetaling)."
    ))

    heading(doc, "CSV-eksportformat", 2, "13.3")
    body(doc, (
        "CSV-filen (semikolon-separeret, UTF-8) indeholder KUN data pr. medarbejder – "
        "topsummeringen 'Total sum for perioden' skrives ikke med. For hver medarbejder skrives "
        "én række pr. dag i perioden (alle 14, også dage uden aktivitet), efterfulgt af én "
        "'Total løn for [navn]'-række:"
    ))
    header_table(doc,
        ["Kolonne", "Indhold"],
        [
            ["Dato",                        "dd-mm-åååå"],
            ["Lønnummer",                   "Medarbejderens employee_number"],
            ["Normal timer … Øvrig overtid","Format T:mm (fx '7:30'), ikke decimaltal"],
            ["Total tid",                   "Decimaltal med komma (fx '7,50')"],
            ["Total i kr. / Beløb",         "Dansk kr-format (fx '1.234,56'), identiske værdier"],
            ["Vognnummer",                  "Aktivitetens vognnummer – overskrives af fraværstypens navn på fraværsdage"],
        ]
    )
    body(doc, (
        "Eksport kræver at den valgte periode er låst (status 'closed') – knappen giver en "
        "fejlbesked og udfører intet, hvis perioden stadig er åben. Administratorer er undtaget "
        "og kan altid eksportere, uanset periodens status. Eksport LÅSER ikke selv perioden "
        "(det gør kun 'Kør løn' under Lønkørsel)."
    ))

    doc.save(OUT / "Teknisk dokumentation.docx")
    print("Teknisk dokumentation.docx gemt.")


# ══════════════════════════════════════════════════════════════════════════════
#  BRUGERVEJLEDNING
# ══════════════════════════════════════════════════════════════════════════════

def build_bruger():
    doc = Document()

    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)

    cover(doc, "Lønsystem", "Brugervejledning", "Juni 2026")

    # ── Introduktion ───────────────────────────────────────────────────────
    heading(doc, "Introduktion", 1, "1")
    body(doc, (
        "Lønsystemet er et webbaseret program der håndterer registrering og lønberegning "
        "for lastbilchauffører hos Poul Schou A/S. Systemet henter automatisk kørselsdata "
        "fra chaufførernes digitale tachografer (.ddd-filer) og giver mulighed for at "
        "godkende, rette og eksportere løndata til Danløn."
    ))
    body(doc, (
        "Systemet tilgås via en almindelig webbrowser – der er ingen installation nødvendig "
        "på den enkelte computer. Åbn browseren og gå til systemets adresse på det lokale netværk."
    ))

    heading(doc, "Tre hovedområder", 2, "1.1")
    header_table(doc,
        ["Område", "Formål"],
        [
            ["Aktiviteter",  "Viser og håndterer alle chaufførers registrerede arbejdsdage i den aktuelle lønperiode."],
            ["Lønkørsel",    "Beregner og eksporterer løn til Danløn samt genererer timesedler og prøvekørsler."],
            ["Medarbejdere", "Stamdata for alle chauffører: lønnummer, overenskomst, normaltimer, afdeling m.m."],
        ]
    )

    # ── 2. Navigation ─────────────────────────────────────────────────────
    heading(doc, "Navigation", 1, "2")

    heading(doc, "Venstre menu", 2, "2.1")
    body(doc, "Menuen er delt i tre grupper. Hvilke punkter du ser afhænger af din rolles rettigheder (se kapitel 11).")
    two_col_table(doc, [
        ["📋 Aktiviteter",    "Viser aktivitetstabellen for den aktuelle lønperiode."],
        ["💰 Lønkørsel",      "Åbner lønkørselspanelet med prøvekørsel, PDF og CSV-eksport. Kræver 'Lønkørsel'-rettigheden."],
        ["📊 Fraværsoversigt","Samlet oversigt over fravær pr. medarbejder/type, med eksport. Kræver 'Fraværsoversigt'-rettigheden."],
        ["🗓️ Vagtplan",       "Ugentlig/periodisk vagtplanlægning pr. medarbejder. Kræver 'Se vagtplan'; redigering kræver desuden 'Redigér egen linje' eller 'Redigér alle linjer'."],
    ])
    body(doc, "")
    two_col_table(doc, [
        ["📥 Importer .ddd",  "Import af tachografdata fra fil eller mappe. Kræver 'Importer .ddd'-rettigheden."],
        ["👤 Medarbejdere",   "Oversigt over og redigering af medarbejderstamdata. Kræver mindst 'Se medarbejdere'."],
        ["🚛 Vognpark",       "Oversigt over og redigering af køretøjer. Kræver mindst 'Se vognpark'."],
        ["💵 Tillæg",         "Administration af individuelle kr/time-tillæg pr. medarbejder. Kræver 'Administrér medarbejdertillæg'."],
    ])
    body(doc, "")
    two_col_table(doc, [
        ["🔑 Brugere",     "Opret/rediger brugere og roller, se hændelseslog (kapitel 11). Kun synlig med 'Brugerstyring'-rettigheden."],
        ["⚙️ Stamdata",    "Konfiguration af systemets masterdata: overenskomsttyper, overtidssatser, tillæg, løntypekoder (inkl. CSV-kolonneopsætning), fraværstyper, CVR-numre, helligdage, disponentgrupper og aftaletyper. Alle typer og koder kan oprettes, redigeres og slettes. Kræver 'Stamdata'-rettigheden."],
    ])

    heading(doc, "Periodenavigation", 2, "2.2")
    body(doc, (
        "Øverst på skærmen vises den aktuelle lønperiode (14 dage). "
        "Brug piletasterne (‹ og ›) til at navigere til forrige eller næste periode. "
        "Du kan også bruge datovælgeren til at hoppe direkte til en specifik dato: "
        "klik på datofeltet for at åbne en kalender med måned- og årstaldropdown."
    ))

    heading(doc, "Statuschips", 2, "2.3")
    header_table(doc,
        ["Chip", "Farve", "Betydning"],
        [
            ["Afventer",     "Rød",   "Antal aktiviteter der venter på godkendelse."],
            ["Godkendt",     "Grøn",  "Antal godkendte aktiviteter."],
            ["Deaktiveret",  "Grå",   "Antal deaktiverede (inaktive) aktiviteter."],
        ]
    )
    body(doc, "Du kan klikke på en chip for hurtigt at filtrere tabellen til den pågældende status.")

    # ── 3. Import af tachografdata ────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Import af tachografdata", 1, "3")
    body(doc, (
        "Tachografdata gemmes på chaufførernes digitale fører-kort som .ddd-filer. "
        "Disse filer importeres i systemet, som automatisk opretter aktiviteter for "
        "alle registrerede arbejdsdage."
    ))

    heading(doc, "Fremgangsmåde", 2, "3.1")
    for i, step in enumerate([
        "Klik på 'Importer .ddd' i den venstre menu.",
        "Vælg 'Vælg filer' for at importere enkeltfiler, eller 'Vælg mappe' for at importere alle .ddd-filer i en mappe.",
        "En Windows-filhåndterings-dialog åbner sig – naviger til filerne og bekræft valget.",
        "Systemet importerer filerne og viser en pop-up med resultatet: enten '✅ Importering succesfuld', "
        "eller en opdelt oversigt over antal importerede, opdaterede og sprunget over – med den konkrete "
        "årsag til hver sprunget-over-gruppe.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")

    note_box(doc,
        "Systemet springer automatisk allerede importerede aktiviteter over "
        "(baseret på medarbejder + starttidspunkt). Det er sikkert at importere "
        "samme mappe flere gange. Pop-up'en skelner mellem 'sprunget over – allerede "
        "importeret' og 'sprunget over – ukendt førerkortnummer' (med de konkrete "
        "kortnumre), så det altid er tydeligt hvorfor en aktivitet ikke kom med.",
        "GODT AT VIDE"
    )
    note_box(doc,
        "Bliver et førerkort læst af flere gange (fx først midt på en vagt og senere igen "
        "efter vagtens afslutning), genkender systemet nu automatisk den mere komplette "
        "udlæsning og UDVIDER den eksisterende aktivitet i stedet for at springe den over "
        "(rettet 2026-07-02). Var aktiviteten allerede godkendt, bliver den automatisk sat "
        "tilbage til 'Afventer', så den udvidede tid skal godkendes igen, før den tæller med "
        "i lønkørslen.",
        "GODT AT VIDE"
    )
    note_box(doc,
        "Alle importkørsler logges under Brugerstyring → Hændelseslog med samme "
        "opsummering, så du kan slå det op igen senere – også selvom pop-up'en er lukket.",
        "GODT AT VIDE"
    )

    heading(doc, "Forudsætninger", 2, "3.2")
    bullet(doc, "Chaufføren skal være oprettet i systemet med korrekt førerkortnummer.")
    bullet(doc, (
        "Førerkortnummeret i systemet skal være de første 14 tegn af kortnummeret: "
        "landekode (2 bogstaver) + 12 cifre (f.eks. DK000000178901) – IKKE de sidste "
        "2 cifre. De sidste 2 cifre er kortets udskiftnings-/fornyelsesindeks og "
        "ændrer sig hver gang kortet fornys, så de er ikke en del af det stabile "
        "kortnummer."
    ))
    bullet(doc, (
        "Stemmer kortnummeret ikke, importeres 0 aktiviteter uden fejlmelding – kun "
        "'Sprunget over'-tallet stiger. Kontrollér og ret kortnummeret i Stamdata "
        "hvis en import ikke giver de forventede aktiviteter."
    ))

    note_box(doc,
        "Starter chaufførens dag med en kort pause (fx et par minutter inden kørslen "
        "begynder), vises denne pause nu som en del af dagens arbejdstid i "
        "aktivitetstabellen. Pausen aflønnes stadig ikke – den trækkes fra som hvil/pause "
        "ligesom systemets øvrige pauser.",
        "GODT AT VIDE"
    )
    note_box(doc,
        "Km-start og km-slut hentes fra førerkortet, som kun gemmer et begrænset antal "
        "køretøjsbrug-poster. Ældre dage vil derfor kunne mangle km-data, selvom resten "
        "af aktiviteten er importeret korrekt.",
        "GODT AT VIDE"
    )

    # ── 4. Aktivitetstabellen ─────────────────────────────────────────────
    heading(doc, "Aktivitetstabellen", 1, "4")
    body(doc, (
        "Aktivitetstabellen er systemets hovedvisning. Den viser alle chauffører "
        "i rækker og de 14 dage i perioden som kolonner. Hver celle indeholder "
        "aktiviteternes start- og sluttidspunkt."
    ))

    heading(doc, "Farvekodning", 2, "4.1")
    header_table(doc,
        ["Farve", "Betydning"],
        [
            ["Rød/mørk",   "Afventer godkendelse (pending)."],
            ["Grøn",       "Godkendt aktivitet."],
            ["Orange",     "Deaktiveret aktivitet."],
            ["Blå/neutral","Manuel aktivitet (ikke fra tachograf) eller korrigeret aktivitet."],
            ["FERIE/FRI",  "Fraværstype – vises med tekst i grøn boks."],
            ["Mørkegrøn kolonne (#056a10)", "Helligdag – søjleoverskriften fremhæves med mørk grøn baggrund. Halvdagshelligdage vises med '½ fra HH:MM'-badge og navn som tooltip."],
        ]
    )

    heading(doc, "Symboler og advarsler", 2, "4.2")
    two_col_table(doc, [
        ["❗ (udråbstegn)",        "Aktiviteten er under 4 timer lang. Kræver kommentar ved godkendelse. Vises kun når status er Afventer eller Deaktiveret – forsvinder ved godkendelse."],
        ["⚠️ (advarselstrekant)", "Aktiviteten er over 12 timer lang. Advarsel om usædvanlig lang vagt. Vises kun når status er Afventer eller Deaktiveret – forsvinder ved godkendelse."],
        ["(K) prefix",             "Aktiviteten er et barn af en opdelt aktivitet."],
    ])

    heading(doc, "Filtrering", 2, "4.3")
    body(doc, "Brug filtrene i toppen af tabellen til at indsnævre visningen:")
    bullet(doc, "Vis alle aktiviteter, kun afventende, kun godkendte eller kun deaktiverede.", "Status: ")
    bullet(doc, "Filtrer til én bestemt chauffør.", "Medarbejder: ")
    bullet(doc, "Vis kun chauffører fra den valgte afdeling. Medarbejder-filteret opdateres automatisk.", "Afdeling: ")

    heading(doc, "Vagter over midnat og periodegrænser", 2, "4.4")
    body(doc, (
        "En kørselsvagt der starter på en søndag eller en helligdag og strækker sig over midnat vises "
        "opdelt i to – ét stykke i hver af de to datokolonner, med sit eget klokkeslæt-interval. Det er "
        "kun en visningsmæssig opdeling (samme princip som lønberegningen bruger for søn-/helligdage); "
        "klik på et af stykkerne åbner altid den samlede, oprindelige aktivitet. Vagter der starter en "
        "almindelig hverdag eller lørdag vises som hidtil (start i én kolonne, slut i den næste)."
    ))
    body(doc, (
        "En vagt der starter sidst i én periode og slutter ind i den næste periode, vises nu i begge "
        "perioders aktivitetstabel, så den ikke 'forsvinder' ved periodeskifte."
    ))
    note_box(doc,
        "Registrerer eller retter du en aktivitet, hvis dato hører til en periode der allerede er "
        "kørt løn på (status 'Lukket'), lægges den automatisk over i den EFTERFØLGENDE periode i "
        "stedet for at blive afvist. Det gælder både manuel oprettelse/rettelse og import af .ddd-filer.",
        "BEMÆRK"
    )

    # ── 5. Aktivitetsdetaljer og godkendelse ──────────────────────────────
    doc.add_page_break()
    heading(doc, "Aktivitetsdetaljer og godkendelse", 1, "5")
    body(doc, (
        "Klik på et tidsemærke i tabellen for at åbne detaljevisningen for aktiviteten. "
        "Her kan du se alle oplysninger og foretage ændringer."
    ))

    heading(doc, "Hvad kan du se?", 2, "5.1")
    two_col_table(doc, [
        ["Start- og sluttid",        "Registreret tidspunkt fra tachografen (eller manuelt angivet)."],
        ["Effektiv tid",              "Samlet arbejdstid fratrukket pauser."],
        ["Aktivitetsfordeling",       "Fordeling mellem kørsel, rådighedstid, andet arbejde og hvil/pause (fra tachografen)."],
        ["Detaljeret tidslinje",      "Alle segmenter med præcise tidspunkter."],
        ["Status og godkender",       "Hvem har godkendt og hvornår."],
    ])

    heading(doc, "Hvad kan du gøre?", 2, "5.2")

    p = doc.add_paragraph()
    r = p.add_run("Ret starttid / sluttid")
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11)
    body(doc, (
        "Angiv en ny dato og et nyt klokkeslæt. Brug piltasterne i time- og minutfeltet "
        "eller skriv tallene direkte. Klik 'Gem rettelse' for at gemme. "
        "Systemet bevarer de originale tachograftider automatisk."
    ), space_after=4)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Fortryd rettelse")
    r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(11)
    body(doc, "Nulstiller tiderne til de originale værdier fra tachografen.", space_after=4)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Opdel aktivitet")
    r3.bold = True; r3.font.name = "Arial"; r3.font.size = Pt(11)
    body(doc, (
        "Del en vagt i to – fx hvis en chauffør skifter tur midt på dagen. "
        "Vælg tidspunktet for opdelingen. Systemet opretter to nye aktiviteter "
        "og fordeler pauser og segmenter korrekt."
    ), space_after=4)

    p4 = doc.add_paragraph()
    r4 = p4.add_run("Godkend")
    r4.bold = True; r4.font.name = "Arial"; r4.font.size = Pt(11)
    body(doc, (
        "Godkend aktiviteten så den indgår i lønkørslen. Angiv dine initialer. "
        "Aktiviteter under 4 timer kræver altid en kommentar."
    ), space_after=4)

    p5 = doc.add_paragraph()
    r5 = p5.add_run("Deaktiver")
    r5.bold = True; r5.font.name = "Arial"; r5.font.size = Pt(11)
    body(doc, (
        "Markér aktiviteten som inaktiv. Den indgår ikke i lønberegningen "
        "men slettes ikke fra databasen."
    ), space_after=8)

    note_box(doc,
        "KUN godkendte aktiviteter medtages i lønkørslen. Husk at gennemgå og "
        "godkende alle relevante aktiviteter inden du kører løn.",
        "VIGTIGT"
    )

    # ── 6. Manuelle aktiviteter ───────────────────────────────────────────
    heading(doc, "Manuelle aktiviteter", 1, "6")
    body(doc, (
        "Brug manuelle aktiviteter til at registrere ferie, afspadsering, fri, "
        "skole/kursus eller andre vagter der ikke fremgår af tachografen."
    ))

    heading(doc, "Oprettelse", 2, "6.1")
    for i, step in enumerate([
        "Klik på '+ Tilføj aktivitet' knappen i aktivitetstabellen.",
        "Vælg medarbejder og aktivitetstype.",
        "Angiv start- og sluttidspunkt.",
        "Tilføj eventuelle pauser ved at klikke '+ Tilføj pause' (se afsnit 6.3).",
        "Udfyld eventuelt turnummer og pålæsning/aflæsning (kun for normale aktiviteter).",
        "Klik 'Gem aktivitet'.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")

    heading(doc, "Fraværstyper og overnatning", 2, "6.2")
    body(doc, "Når du vælger en fraværstype eller overnatning, sker følgende automatisk:")
    bullet(doc, "Felterne for turnummer, pålæsning og aflæsning skjules (ikke relevante).")
    bullet(doc, "Aktiviteten godkendes automatisk (approved_by sættes til din bruger) – kræver ikke separat godkendelse bagefter.")
    bullet(doc, "Ferie, sygdom, feriefri m.fl.: starttidspunktet sættes automatisk til 06:00, og sluttidspunktet beregnes ud fra medarbejderens normaltimer den pågældende dag. Vælger du 'Til dato' for at oprette en periode, oprettes én aktivitet PR. HVERDAG i perioden (ikke én sammenhængende aktivitet) – hver dag tæller sine egne normaltimer (typisk 7,4 t).")
    bullet(doc, "Afspadsering som periode ('Til dato' udfyldt) følger samme regel: 7,4 t (eller medarbejderens skemalagte timer) pr. hverdag, uanset klokketid. En enkelt afspadseringsdag (uden 'Til dato') kan derimod redigeres til en delvis dag med selvvalgt start-/sluttid, og den faktiske varighed bruges da i lønberegningen.")

    body(doc, (
        "De tilgængelige fraværstyper administreres i Stamdata-modulet under fanen 'Fraværstyper'. "
        "Her kan du aktivere og deaktivere typer, tilføje nye typer og slette eksisterende."
    ))
    header_table(doc,
        ["Type", "Beskrivelse"],
        [
            ["Ferie",          "Registrerer en feriedag. Start: 06:00. Slut: 06:00 + normaltimer for dagen. Tælles i timeoversigten, men er som default IKKE med i Danløn-CSV'en (kan slås til i Stamdata)."],
            ["Afspadsering",   "Registrerer afspadsering. Som periode ('Til dato') tæller hver hverdag 7,4 t/skemalagte timer; som enkeltdag bruges den faktiske start-/sluttid."],
            ["Fri",            "Registrerer fridag."],
            ["Skole/kursus",   "Registrerer skole- eller kursusdag."],
            ["Overnatning",    "Registrerer en overnatning (flat sats pr. forekomst – ikke timer). Angiv blot datoen. Satsen hentes automatisk fra Stamdata (Tillæg-fanen)."],
        ]
    )

    heading(doc, "Registrering af pauser", 2, "6.3")
    body(doc, (
        "For normale aktiviteter kan du registrere de pauser chaufforen har holdt i løbet af "
        "arbejdsdagen. Pauserne fratrækkes automatisk i den effektive arbejdstid og indgår korrekt "
        "i beregningen af overtidstillæg."
    ))
    for i, step in enumerate([
        "Angiv aktivitetens starttidspunkt (påkrævet – pausedatoen arves herfra).",
        "Klik '+ Tilføj pause'. En dialogboks åbner sig med titel 'Pause 1', 'Pause 2' osv.",
        "Angiv pausens starttidspunkt og sluttidspunkt (kun klokkeslæt – dato sættes automatisk).",
        "Klik 'Tilføj'. Pausen vises i listen med start- og sluttid og en ×-slet-knap.",
        "Gentag for yderligere pauser. Klik × for at fjerne en pause.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")
    note_box(doc,
        "Pausefelterne vises kun for normale aktiviteter. "
        "For fraværstyper (ferie, fri m.fl.) og overnatning er pausesektionen skjult.",
        "GODT AT VIDE"
    )
    body(doc, "Pauserne påvirker tre steder i systemet:")
    bullet(doc, "Sum, effektiv tid – vises i aktivitetsdetaljerne og er fratrukket pausetid.")
    bullet(doc, "Aktivitetsfordeling – manuelle aktiviteter med pauser viser 'Effektiv tid' (blå) og 'Hvil/pause' (grå) i %-bjælken.")
    bullet(doc, "Lønberegning – pauser fratrækkes i det præcise tidsvindue de afholdes, så overtidstillæg beregnes korrekt.")

    # ── 7. Medarbejdere ───────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Medarbejdere", 1, "7")
    body(doc, "Klik på 'Medarbejdere' i venstre menu for at se og redigere medarbejderstamdata.")

    heading(doc, "Opret ny medarbejder", 2, "7.1")
    body(doc, "Klik '+ Opret medarbejder' og udfyld felterne:")

    header_table(doc,
        ["Felt", "Påkrævet", "Beskrivelse"],
        [
            ["Aftale",              "Ja", "Vælg fra listen der stammer fra Stamdata (Aftale-fanen) – som udgangspunkt 'Timelønnet, fast arbejdstid' eller 'Timelønnet, ikke fastlagt arbejdstid'. Administratorer kan tilføje flere aftaletyper i Stamdata."],
            ["Overenskomsttype",    "Afhænger af Aftale", "Bestemmer timesatsen. Vælg fra listen der stammer fra Stamdata (Overenskomsttyper-fanen). Feltet er kun obligatorisk (rød *) hvis den valgte Aftale-type kræver det – styres pr. aftaletype i Stamdata."],
            ["Disponentgrupper",    "Nej","Afdeling(er) – afkrydsningsbokse, en medarbejder kan tilhøre flere grupper samtidig. Bruges til at filtrere aktivitetstabellen og fraværsoversigt-eksporten."],
            ["Lønnummer",           "Ja", "Unikt lønnummer (bruges i Danløn-eksporten)."],
            ["Førerkortnummer",     "Nej","EU-førerkortnummer – påkrævet for at importere .ddd-filer korrekt."],
            ["Navn",                "Ja", "Fornavn og efternavn."],
            ["Adresse / postnummer","Nej","Adresseoplysninger."],
            ["E-mail / telefon",    "Nej","Kontaktoplysninger."],
            ["Ansættelsesdato",     "Ja", "Startdato for ansættelsen. Klik på feltet for at åbne en kalender med måned- og årstaldropdown."],
            ["Fratrædelsesdato",    "Nej","Udfyldes KUN ved fratrædelse. Default er 31-12-9999. Klik på feltet for at åbne kalender-pickeren."],
            ["Aktiv",               "—", "Afkryd for at medarbejderen er aktiv i systemet."],
            ["Fuldlønnet",          "—", "Afkryd hvis medarbejderen er fuldlønnet."],
        ]
    )
    note_box(doc,
        "De tilgængelige aftaletyper administreres i Stamdata-modulet under fanen 'Aftale'. Her "
        "kan en administrator ændre teksten på de to oprindelige typer, samt tilføje flere "
        "aftaletyper og bestemme om Overenskomsttype skal være påkrævet for hver af dem. "
        "Vælges en type der ikke kræver Overenskomsttype, forsvinder den røde stjerne ved feltet, "
        "og det kan stå tomt.",
        "GODT AT VIDE"
    )
    note_box(doc,
        "Hvis navn (for- og efternavn) eller førerkortnummer allerede findes på en anden medarbejder "
        "(aktiv eller inaktiv), viser systemet en advarsel ved oprettelse. Ved navnesammenfald kan du "
        "vælge 'OK, opret alligevel' for at oprette begge, eller 'Ændre' for at vende tilbage og rette. "
        "Ved sammenfald på førerkortnummer kan du KUN vælge 'Ændre' – to medarbejdere kan ikke dele samme "
        "førerkortnummer i databasen.",
        "BEMÆRK"
    )

    heading(doc, "Timefordeling", 2, "7.2")
    body(doc, (
        "Timefordelingen angiver medarbejderens normaltimer pr. ugedag i en 14-dages periode "
        "(uge A og uge B, svarende til ulige og lige ISO-uger). Standardværdier er:"
    ))
    bullet(doc, "Mandag–torsdag: 7,5 timer")
    bullet(doc, "Fredag: 7 timer")
    bullet(doc, "Lørdag–søndag: 0 timer")
    body(doc, (
        "For hver dag kan du udfylde ENTEN et timeantal ELLER en fra/til-tid – begge felter står altid "
        "synlige side om side. Udfylder du fra/til-tid, beregner systemet automatisk timeantallet "
        "(sluttid minus starttid – en sluttid tidligere end starttid tolkes som en vagt der går over "
        "midnat) og overskriver timeantal-feltet. Et 'Ugentligt timeantal'-felt nederst i tabellen "
        "summerer automatisk timerne for hver uge, og opdateres med det samme uanset hvilken af de to "
        "indtastningsmåder du bruger."
    ))
    body(doc, (
        "Normaltimerne bruges til beregning af overtidstillæg og til at auto-udfylde sluttidspunktet ved registrering af ferie."
    ))

    heading(doc, "Anciennitetsvarsler", 2, "7.3")
    body(doc, (
        "Systemet kontrollerer løbende om medarbejdere har nået 9 måneders anciennitet "
        "og en tilsvarende overenskomsttype med højere sats. Hvis det er tilfældet, vises "
        "et pop-up-varsel. Varslet er styret af tilladelsen 'Anciennitetsvarsel' i Brugerstyring – "
        "som standard er den slået til for Administrator og Lønbogholder og fra for Disponent og Kontor. "
        "En administrator kan klikke tilladelsen til eller fra for enhver rolle. "
        "Husk at opdatere overenskomsttypen for medarbejderen."
    ))
    body(doc, (
        "Når du klikker \"Ændring foretaget\" i varslet, gemmes afvisningen på medarbejderen i databasen "
        "(feltet anciennitet_dismissed_at). Varslet vises ikke igen – uanset computer eller browser. "
        "Hvis overenskomsttypen ændres på medarbejderen, nulstilles afvisningen automatisk, "
        "så varslet kan dukke op igen hvis ancienniteten stadig er relevant."
    ))

    heading(doc, "Tillæg", 2, "7.4")
    body(doc, (
        "Ud over overenskomsttypens timesats kan en medarbejder have et individuelt kr/time-tillæg. "
        "Klik på 'Tillæg' i venstre menu for at administrere dette (kræver rettigheden 'Administrér "
        "medarbejdertillæg')."
    ))
    for i, step in enumerate([
        "Søg medarbejderen frem i søgefeltet (navn eller lønnummer).",
        "Klik på medarbejderen. En boks åbner sig med medarbejderens fulde tillægshistorik: status "
        "(Aktiv/Inaktiv), lønnummer, tillægsnavn, type, gyldighedsperiode og værdi i kr. Brug 'Fra'/'Til'-"
        "felterne til at indsnævre visningen til en bestemt periode – lader du dem stå tomme, vises hele "
        "historikken.",
        "Klik '+ Tilføj' (enten i boksen eller i toolbaren øverst på siden) for at oprette et nyt tillæg: "
        "vælg medarbejder (forudvalgt hvis du kom fra dennes boks), startdato (default dags dato) og "
        "værdi i kr/time. Tillægsnavn og type sættes automatisk og kan ikke ændres.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")
    note_box(doc,
        "Når et nyt tillæg oprettes, bliver medarbejderens tidligere tillæg automatisk gjort inaktivt "
        "fra og med dagen før den nye startdato. Der findes ikke en 'rediger'-funktion – en fejl rettes "
        "ved at oprette et nyt tillæg med den korrekte værdi.",
        "GODT AT VIDE"
    )
    body(doc, (
        "Har en medarbejder et aktivt tillæg der ikke længere skal gælde, klik 'Afslut' på den aktive "
        "række. Tillægget fortsætter med at gælde resten af den igangværende lønperiode og bortfalder "
        "først fra den efterfølgende periode – det stopper altså ikke med øjeblikkelig virkning midt i "
        "en periode."
    ))
    note_box(doc,
        "Tillægget lægges automatisk til medarbejderens grundsats i selve lønberegningen – både i "
        "prøvekørslen, PDF-timesedlerne, Danløn CSV'en (kode 1/normaltid) og i afspadsering, sygdom, "
        "barsel og skole/kursus. Det kræver ingen særskilt handling ud over at oprette/afslutte "
        "tillægget her.",
        "VIGTIGT"
    )
    body(doc, (
        "Findes der et aktivt tillæg på en medarbejder, vises værdien også som et read-only felt under "
        "Overenskomsttype, når du åbner medarbejderen under fanen 'Medarbejdere'. Feltet kan ikke "
        "redigeres derfra – kun under fanen 'Tillæg'."
    ))

    # ── 8. Helligdagskalender ─────────────────────────────────────────────
    heading(doc, "Helligdagskalender", 1, "8")
    body(doc, (
        "Systemet vedligeholder automatisk en kalender over danske helligdage. "
        "Helligdage fremhæves i aktivitetskalenderen og danner grundlag for fremtidig "
        "beregning af helligdagstillæg."
    ))

    heading(doc, "Markering i aktivitetskalenderen", 2, "8.1")
    body(doc, (
        "Dage der er helligdage fremhæves med mørk grøn baggrund (#056a10) i "
        "aktivitetskalenderens kolonneoverskrifter. Halvdagshelligdage (1. maj og Grundlovsdag) "
        "vises med et '½ fra 12:00'-badge. Hold musen over overskriften for at se helligdagens navn."
    ))
    note_box(doc,
        "Helligdagskalenderen genereres automatisk for de næste 5 år ved serveropstart. "
        "Du behøver ikke gøre noget – den opdaterer sig selv.",
        "GODT AT VIDE"
    )

    heading(doc, "Administration af helligdage (Stamdata → Helligdage)", 2, "8.2")
    body(doc, (
        "Brugere med rettigheden 'Administrér helligdage' kan administrere helligdagskalenderen "
        "via Stamdata-menupunktet → fanen 'Helligdage':"
    ))
    header_table(doc,
        ["Handling", "Beskrivelse"],
        [
            ["Generer for år",   "Klik 'Generer [årstal]' for at auto-generere alle danske helligdage for det valgte år. Eksisterende datoer overskrives ikke."],
            ["Tilføj manuelt",   "Klik '+Tilføj helligdag' for at oprette en særlig fridag med valgfri halvdagstid (fx en lukkedag)."],
            ["Slet helligdag",   "Klik 'Slet' på en helligdag for at fjerne den fra kalenderen. Auto-genererede og manuelle helligdage kan begge slettes."],
        ]
    )
    two_col_table(doc, [
        ["Faste helligdage",      "Nytårsdag, 1. maj (½), Grundlovsdag (½), Juleaftensdag, 1. og 2. juledag, Nytårsaftensdag"],
        ["Bevægelige helligdage", "Skærtorsdag, Langfredag, Påskedag, 2. påskedag, Kristi Himmelfartsdag, Pinsedag, 2. pinsedag"],
        ["Store Bededag",         "Medtages ikke – afskaffet fra 2024"],
    ])

    # ── 9. Lønkørsel ─────────────────────────────────────────────────────
    heading(doc, "Lønkørsel", 1, "9")
    body(doc, (
        "Klik på 'Lønkørsel' i venstre menu for at beregne og eksportere løn. "
        "Husk: kun aktiviteter med status 'Godkendt' indgår i beregningen."
    ))

    heading(doc, "Prøvekørsel", 2, "9.1")
    body(doc, (
        "Prøvekørslen beregner lønnen og gemmer resultatet som en Excel-fil. "
        "Brug denne til at tjekke tallene inden den endelige kørsel."
    ))
    bullet(doc, "Klik 'Prøvekørsel'.")
    bullet(doc, "Vælg evt. en bestemt medarbejder, eller lad feltet stå tomt for alle.")
    bullet(doc, "Klik 'Gennemse' for at vælge hvilken mappe filen skal gemmes i (foreslår Downloads-mappen).")
    bullet(doc, "Klik 'Kør prøvekørsel'. Filen gemmes og pop-up-vinduet lukker automatisk.")

    body(doc, "Excel-filen viser alle 14 dage i perioden for hver medarbejder:")
    bullet(doc, "Dage med normal aktivitet: timer fordelt på Normal tid, Tillæg 05-06, Tillæg OT1-3 og Øvrigt OT.")
    bullet(doc, "Fraværsdage (Ferie, Fri, Afspadsering, Skole/kursus): vises med fraværstype i Normal tid-kolonnen.")
    bullet(doc, "Dage uden registrering: vises med 0 på alle kolonner.")
    bullet(doc, "Første linje pr. medarbejder er markeret med grøn baggrund. Tom linje adskiller medarbejdere.")

    heading(doc, "Dan PDF'er", 2, "9.2")
    body(doc, "Genererer individuelle timesedler til alle medarbejdere som PDF-filer.")
    bullet(doc, "Klik 'Dan PDF'er'.")
    bullet(doc, "Vælg mappe at gemme i (brug 'Gennemse').")
    bullet(doc, "Klik 'Dan PDF'er'. Én PDF pr. medarbejder gemmes i mappen.")

    heading(doc, "Kør løn (Danløn CSV)", 2, "9.3")
    body(doc, (
        "Eksporterer den endelige lønfil til Danløn. "
        "Filen er i CSV-format og kan importeres direkte i Danløn."
    ))
    bullet(doc, "Klik 'Kør løn'.")
    bullet(doc, "Bekræft at du er klar til at køre endelig løn.")
    bullet(doc, "CSV-filen downloades til din computer.")
    body(doc, (
        "CSV-filen indeholder op til 6 kolonner per lønpost: CVR-nummer, medarbejdernr., lønkode, "
        "antal (timer eller forekomster), sats og evt. total. "
        "Hvilke kolonner der medtages afhænger af opsætningen i Stamdata → Løntypekoder."
    ))
    body(doc, (
        "'Kør løn'-knappen er grå og kan ikke bruges, hvis der enten er aktiviteter i perioden, "
        "der endnu ikke er godkendt eller deaktiveret, eller hvis lønperioden allerede er kørt og låst. "
        "I begge tilfælde vises en advarsel, hvis man alligevel klikker på knappen. Afventende aktiviteter "
        "skal godkendes eller deaktiveres under fanen 'Aktiviteter', før knappen aktiveres igen. En allerede "
        "låst periode kan genåbnes under 'Administration' (kræver rettigheden 'Åbn låst lønperiode'), hvis "
        "der skal foretages ændringer og køres løn igen."
    ))

    heading(doc, "CSV-kolonneopsætning per løntypekode", 2, "9.3.1")
    body(doc, (
        "Under Stamdata → Løntypekoder kan du for hver løntype styre præcis hvad der eksporteres i CSV-filen:"
    ))
    header_table(doc,
        ["Indstilling", "Valgmuligheder", "Beskrivelse"],
        [
            ["Antal-type",    "Timer / Antal",                 "Timer: eksporterer timetal med 2 decimaler. Antal: eksporterer antal forekomster som heltal (fx overnatninger)."],
            ["Sats-kilde",    "Timesats, OT-satser, Salt, Overnatning, Dagpenge §56", "Bestemmer hvilken sats der bruges til kolonnen Sats og beregning af Total."],
            ["Inkluder sats", "Ja / Nej",                      "Ja: sats-kolonnen skrives i CSV. Nej: satsen udelades fra CSV, men bruges stadig til beregning af Total."],
            ["Inkluder total","Ja / Nej",                      "Ja: en ekstra total-kolonne tilføjes (antal × sats). Nej: ingen total-kolonne."],
        ]
    )
    note_box(doc,
        "Inkluder sats = Nej og Inkluder total = Ja giver en CSV med kun antal og total – "
        "nyttigt når Danløn selv beregner satsen ud fra lønkoden.",
        "TIP"
    )

    heading(doc, "Hvad beregnes?", 2, "9.4")
    body(doc, "For hver godkendt aktivitet beregnes følgende lønposter:")
    header_table(doc,
        ["Lønpost", "Beskrivelse"],
        [
            ["Normal løn",         "Timesats × normaltimer for dagen (fra medarbejderens timefordeling)."],
            ["Tillæg 05-06",       "Forhøjet sats for arbejde mellem kl. 05:00 og 06:00."],
            ["Tillæg 18-21 / OT",  "Forhøjet sats for arbejde kl. 18-21 og overarbejde op til 3 timer."],
            ["Nat / ekstra OT",    "Højeste sats for natarbejde (21-05) og overarbejde over 3 timer."],
            ["Pålæsning",          "Tillæg for pålæsningstid (minutter × minutsats)."],
            ["Aflæsning",          "Tillæg for aflæsningstid (minutter × minutsats)."],
            ["Salttillæg",         "Tillæg pr. time for kørsel med salt (registreret på aktiviteten). Sats fra Stamdata (Tillæg-fanen)."],
            ["Overnatning",        "Fast sats pr. overnatning (antal forekomster × sats). Sats fra Stamdata (Tillæg-fanen)."],
        ]
    )

    heading(doc, "Lørdage, søndage og helligdage", 2, "9.5")
    body(doc, (
        "Lørdage, søndage og helligdage har deres egne lønkoder ud over de tre almindelige "
        "tillæg (kode 1, 7, 8 og 9). Du behøver ikke gøre noget særligt ved registreringen – "
        "systemet genkender automatisk dagtypen ud fra helligdagskalenderen og ugedagen."
    ))
    header_table(doc,
        ["Dag", "Hvad sker der?"],
        [
            ["Lørdag",   "Regnes som en almindelig hverdag. Har medarbejderen ingen garanterede timer den lørdag (det typiske), bliver kørte timer automatisk til kode 8/9 efter samme regler som en hverdags overarbejde."],
            ["Søndag / helligdag", "Garanterede timer udbetales som kode 4 (fuldlønnet) eller kode 63 (timelønnet), oveni evt. kørsel. Alt kørt arbejde går til kode 1 + kode 9 – uanset klokkeslæt."],
            ["1. maj / Grundlovsdag", "Halve helligdage: fri fra kl. 12:00. Kørsel før 12 tæller normalt, kørsel efter 12 følger særlige regler for disse dage."],
        ]
    )
    note_box(doc,
        "En vagt der starter aftenen/natten før en søndag eller helligdag og fortsætter ind i "
        "den, vil i prøvekørslen og PDF-timesedlen kunne vise sig som to linjer i stedet for én "
        "– én for delen inden midnat, én for delen efter. Det er korrekt: de to dele følger hver "
        "sin dags regler. En vagt der fortsætter fra en almindelig hverdag/lørdag ind i en anden "
        "almindelig hverdag/lørdag vises derimod stadig som én sammenhængende linje.",
        "GODT AT VIDE"
    )

    # ── 10. Vigtige regler ────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Vigtige regler og opmærksomhedspunkter", 1, "10")

    heading(doc, "Inden lønkørsel – tjekliste", 2, "10.1")
    for item in [
        "Alle .ddd-filer er importeret for perioden.",
        "Alle relevante aktiviteter er gennemgået og godkendt.",
        "Aktiviteter med ! (under 4 timer) er godkendt med kommentar.",
        "Eventuelle rettelser er kontrolleret – brug 'Fortryd rettelse' ved fejl.",
        "Prøvekørsel er gennemgået og tallene ser korrekte ud.",
    ]:
        bullet(doc, item)

    heading(doc, "Opdatering af satser", 2, "10.2")
    body(doc, (
        "Timesatser og overtidssatser opdateres via Stamdata-modulet i systemets brugerflade. "
        "Ændringerne slår igennem ved næste beregning uden genstart."
    ))
    for i, step in enumerate([
        "Klik på '⚙️ Stamdata' i venstre menu (kræver administrator-login).",
        "Vælg den relevante fane: 'Overenskomsttyper', 'Overtidssatser' eller 'Tillæg'.",
        "Klik på 'Rediger' ud for den sats der skal ændres.",
        "Indtast den nye sats og klik 'Gem'.",
    ], 1):
        bullet(doc, step, f"Trin {i}: ")
    two_col_table(doc, [
        ["Overenskomsttyper-fanen", "Timesatser for overenskomsttyper (normal løn)."],
        ["Overtidssatser-fanen",    "Satser for de tre overtidstillæg (1 time før, 1-3 timer efter, øvrigt)."],
        ["Tillæg-fanen",            "Sats for salttillæg (pr. time), overnatning (pr. forekomst) og §56."],
    ])
    note_box(doc,
        "Rediger ikke Excel-filerne med forventning om at ændringerne slår igennem. "
        "Excel bruges KUN til den initiale seeding af databasen ved første opstart. "
        "Al efterfølgende konfiguration sker via Stamdata-modulet.",
        "VIGTIGT"
    )

    heading(doc, "Fejl og usædvanlige situationer", 2, "10.3")
    header_table(doc,
        ["Situation", "Hvad gør du?"],
        [
            ["Forkert start- eller sluttid importeret",
             "Åbn aktiviteten og brug 'Ret starttid/sluttid'. Brug 'Fortryd rettelse' hvis du fortryder."],
            ["Chauffør er ikke i systemet (import fejler)",
             "Opret medarbejderen med korrekt førerkortnummer og importér .ddd-filen igen."],
            ["Import viser 0 importerede og mange 'sprunget over', selvom filerne er nye",
             "Kontrollér medarbejderens førerkortnummer i Stamdata: det skal være de "
             "første 14 tegn (landekode + 12 cifre), uden de sidste 2 udskiftnings-/"
             "fornyelsescifre. Ret nummeret og importér filen igen."],
            ["Aktivitet skal deles i to (f.eks. skift af tur)",
             "Brug 'Opdel aktivitet' og angiv tidspunktet for opdelingen."],
            ["Aktivitet skal ignoreres",
             "Brug 'Deaktiver' – aktiviteten medtages ikke i lønberegningen men slettes ikke."],
            ["Medarbejder har nået 9 måneder",
             "Systemet viser et varsel. Opdater overenskomsttypen for medarbejderen."],
            ["Lønkørslen viser ikke nye lønposter (fx overnatning)",
             "Genstart serveren: stop og start igen via Preview-panelet i Claude Code. Systemet indlæser ny beregningslogik ved opstart."],
        ]
    )

    note_box(doc,
        "Fratrædelsesdatoen for aktive medarbejdere skal stå på 31-12-9999. "
        "Udfyld kun fratrædelsesdatoen den dag en medarbejder forlader virksomheden.",
        "HUSK"
    )

    # ── 11. Brugerstyring og roller ───────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Brugerstyring og roller", 1, "11")
    body(doc, (
        "Klik på '🔑 Brugere' i venstre menu for at oprette login-brugere, styre hvilke "
        "roller de har, og se hændelsesloggen. Menupunktet er kun synligt for brugere hvis "
        "rolle har rettigheden 'Brugerstyring'. Siden har tre faner: Brugere, Roller og "
        "Hændelseslog."
    ))

    heading(doc, "Brugere-fanen", 2, "11.1")
    body(doc, "Klik '+ Opret bruger' og udfyld navn, initialer/brugernavn (bruges til login), mail (valgfri), rolle og adgangskode.")
    two_col_table(doc, [
        ["Rediger",                "Ændrer navn, initialer, mail, rolle eller adgangskode på en eksisterende bruger."],
        ["Aktiver / Deaktiver",    "Slår login til/fra uden at slette brugeren eller dens historik. En deaktiveret bruger kan ikke logge ind."],
        ["Din egen konto",         "Kan redigeres, men ikke deaktiveres – knappen erstattes af teksten '(dig selv)'."],
    ])
    note_box(doc,
        "Initialer skal være unikke (store/små bogstaver behandles ens) og bruges som "
        "brugernavn ved login.",
        "BEMÆRK"
    )

    heading(doc, "Roller-fanen", 2, "11.2")
    body(doc, (
        "En rolle er en navngivet samling af rettigheder. Hver bruger har præcis én rolle "
        "(sat i Brugere-fanen), og rollen bestemmer hvad brugeren kan se og gøre i systemet."
    ))
    header_table(doc,
        ["Rolle", "Type", "Typiske rettigheder ved oprettelse"],
        [
            ["Administrator", "Systemrolle – kan ikke slettes, rettigheder kan ikke redigeres",
             "Alt: lønkørsel, import, brugerstyring, åbn låst periode, administrér baselines, godkend aktiviteter, se kalender, sæt springertillæg, vagtplan (se + redigér alle linjer)"],
            ["Lønbogholder", "Almindelig rolle – kan redigeres/slettes",
             "Lønkørsel, fraværsoversigt, import, se/tilføj medarbejdere, se/tilføj vogne, anciennitetsvarsel, godkend aktiviteter, se kalender, administrér medarbejdertillæg, sæt springertillæg, vagtplan (se + redigér alle linjer)"],
            ["Disponent", "Almindelig rolle – kan redigeres/slettes",
             "Se medarbejdere, se vogne, godkend aktiviteter, se kalender, sæt springertillæg, vagtplan (se + redigér alle linjer)"],
            ["Kontor", "Almindelig rolle – kan redigeres/slettes",
             "Se vagtplan, redigér egen linje i vagtplan"],
        ]
    )
    note_box(doc,
        "Kun 'Administrator' er en beskyttet systemrolle. De øvrige roller – herunder "
        "'Lønbogholder' og 'Disponent' – kan en administrator frit redigere (rettigheder og "
        "visningsnavn) eller slette, og der kan oprettes helt nye, skræddersyede roller efter "
        "behov. 'Kontor' er et eksempel på en sådan senere oprettet rolle, målrettet "
        "medarbejdere der kun skal kunne se og redigere deres egen linje i Vagtplan.",
        "BEMÆRK"
    )
    body(doc, "Klik '+ Opret rolle' for at oprette en ny rolle: angiv et internt rollenavn (kun bogstaver, tal og underscore), et visningsnavn, og afkryds de ønskede rettigheder fra listen nedenfor. Klik 'Rediger' på en eksisterende, ikke-system-rolle for at ændre visningsnavn og rettigheder. En rolle kan kun slettes, hvis ingen brugere har den.")

    heading(doc, "Rettighedsoversigt", 2, "11.3")
    header_table(doc,
        ["Rettighed (nøgle)", "Vises som", "Giver adgang til"],
        [
            ["payroll",                     "Lønkørsel",                          "Menupunktet 'Lønkørsel': prøvekørsel, PDF-timesedler og Danløn CSV-eksport."],
            ["absence_overview",            "Fraværsoversigt",                    "Menupunktet 'Fraværsoversigt'."],
            ["import_ddd",                  "Importer .ddd",                      "Menupunktet 'Importer .ddd' – import af tachografdata."],
            ["user_management",             "Brugerstyring",                     "Menupunktet '🔑 Brugere' – denne side (brugere, roller, hændelseslog)."],
            ["reopen_period",               "Åbn låst lønperiode",                "Mulighed for at genåbne en periode hvor der allerede er kørt løn."],
            ["stamdata",                     "Stamdata",                          "Menupunktet '⚙️ Stamdata' – alle ni faner (se kapitel 7)."],
            ["view_employees",              "Se medarbejdere",                   "Menupunktet 'Medarbejdere' (læse-adgang)."],
            ["manage_employees",            "Tilføj medarbejdere",               "Opret/rediger medarbejdere under 'Medarbejdere'."],
            ["view_vehicles",               "Se vognpark",                       "Menupunktet 'Vognpark' (læse-adgang)."],
            ["manage_vehicles",             "Tilføj vogn",                       "Opret/rediger køretøjer under 'Vognpark'."],
            ["manage_employee_supplements", "Administrér medarbejdertillæg",     "Menupunktet 'Tillæg' – oprette/afslutte kr/time-tillæg."],
            ["manage_holidays",             "Administrér helligdage",            "Stamdata → Helligdage: opret, auto-generer og slet."],
            ["anciennitet_alert",           "Anciennitetsvarsel",                "Modtag pop-up-varsler om medarbejdere med 9 måneders anciennitet."],
            ["approve_activities",          "Godkend aktiviteter",               "Godkend/deaktiver/ret/opdel aktiviteter i aktivitetstabellen."],
            ["view_calendar",               "Se aktivitetskalender",             "Se aktivitetstabellen (kalendervisningen på forsiden)."],
            ["toggle_springer",             "Sæt springertillæg",                "Afkryds springertillæg-fluebenet i aktivitetsoversigten."],
            ["vagtplan_view",               "Se vagtplan",                       "Menupunktet 'Vagtplan' (læse-adgang)."],
            ["vagtplan_edit_own",           "Redigér egen linje i vagtplan",     "Ret kun sin egen række i Vagtplan (matches på initialer)."],
            ["vagtplan_edit_all",           "Redigér alle linjer i vagtplan",    "Ret alle medarbejderes rækker i Vagtplan."],
            ["payroll_settlement_view",     "Lønafregning (se)",                 "Menupunktet 'Lønafregning': periodetotaler og pr.-medarbejder oversigt (se kapitel 12)."],
            ["payroll_settlement_export",   "Lønafregning (eksport)",            "Eksportér Lønafregning som CSV. Kræver låst periode, medmindre brugeren er administrator."],
        ],
        [Cm(3.5), Cm(4), Cm(8.5)]
    )
    note_box(doc,
        "Rettigheder er additive og uafhængige af hinanden – der er ingen indbygget "
        "hierarki. En rolle med kun 'Se vagtplan' + 'Redigér egen linje i vagtplan' (som "
        "'Kontor') ser IKKE Aktiviteter, Lønkørsel eller andre menupunkter, uanset hvor "
        "'lille' rettigheden lyder.",
        "TEKNISK NOTE"
    )

    heading(doc, "Hændelseslog-fanen", 2, "11.4")
    body(doc, (
        "Viser de seneste 500 hændelser i systemet, nyeste øverst: oprettelse/redigering af "
        "brugere og roller, rolleskift, samt Stamdata-ændringer (opret/redigér/slet). Hver "
        "linje viser tidspunkt, hvilken bruger der udførte handlingen, selve handlingen og en "
        "kort detaljetekst. Søgefeltet filtrerer på tværs af bruger, handling og detaljer; "
        "'⟲ Opdater' genindlæser listen."
    ))

    # ── 12. Lønafregning ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Lønafregning", 1, "12")
    body(doc, (
        "Klik på '🧾 Lønafregning' i venstre menu, lige under 'Lønkørsel'. Fanen giver et samlet "
        "overblik over periodens løn: én total-tabel øverst for hele perioden, og derunder én "
        "tabel pr. medarbejder med alle 14 dage. Menupunktet kræver rettigheden "
        "'Lønafregning (se)', og eksportknappen kræver derudover 'Lønafregning (eksport)' "
        "(se afsnit 11.3)."
    ))

    heading(doc, "Periodetotalen øverst", 2, "12.1")
    body(doc, (
        "Den øverste tabel 'Total sum for perioden' viser: Grundtimeløn inkl. tillæg, Overtid "
        "Timen før, Overtid 1-3 time efter og Øvrig overtid, efterfulgt af delsummen "
        "'Total uden fravær'. Herunder følger op til 11 fraværstyper (Sygdom, Sygdom u. 8 uger, "
        "Barn 1.sygedag, Barn 1.sygedag u. 8 uger, Graviditetsbetinget sygdom, §56 syg, Barsel, "
        "Feriefri, Ferie, Skole/kursus og Afspadsering) og til sidst 'Total sum for denne "
        "periode'. En linje vises KUN hvis den har et beløb større end 0 kr for den viste "
        "periode – har ingen medarbejder fx haft barsel, udelades 'Barsel'-linjen helt."
    ))

    heading(doc, "Medarbejder-tabellerne", 2, "12.2")
    body(doc, (
        "Hver medarbejders kort viser i overskriften: overenskomsttype med sats i kr/time, "
        "personligt tillæg (hvis medarbejderen har et) og springertillæg (hvis aktiveret for "
        "perioden) – alle tre vist som separate tal, ikke lagt sammen. Derunder følger en "
        "tabel med én række pr. dag i perioden (også dage uden aktivitet, som vises med 0)."
    ))
    header_table(doc,
        ["Kolonne", "Indhold"],
        [
            ["Normal timer / Overtid 1 time før / Overtid 1-3 timer efter / Øvrig overtid",
             "Vist som timer:minutter (fx '7:30'), ikke decimaltal."],
            ["Total tid",     "Dagens samlede tid som decimaltal (fx '7,50')."],
            ["Total i kr. / Beløb", "Dagens samlede løn i kr. På en fraværsdag med en betalt "
             "fraværstype (se afsnit 12.1) vises fraværstimernes egen værdi her i stedet for 0,00."],
            ["Vognnummer",    "Aktivitetens vognnummer – på en fraværsdag vises i stedet "
             "fraværstypens navn (fx 'Sygdom'), selvom aktiviteten har et udfyldt vognnummer-felt."],
        ]
    )
    body(doc, (
        "Nederst i hver medarbejders tabel står 'Total løn for [navn]' – den fulde sum for "
        "medarbejderen i perioden: arbejdstid, overtid, salttillæg, springertillæg OG al "
        "fraværsbetaling med et beregnet beløb (se afsnit 12.1 for hvilke typer det gælder)."
    ))

    heading(doc, "Periode-visning", 2, "12.3")
    body(doc, (
        "Lønafregning har ingen egen periodevælger – siden følger i stedet den periode du "
        "senest har navigeret til med pilene (‹ ›) i Aktivitetsoversigten, ligesom Lønkørsel "
        "allerede gør. Har du ikke navigeret til en bestemt periode, vises dagens periode. "
        "Skal du fx tjekke eller eksportere en periode der er kørt løn for og siden er blevet "
        "'gammel' (dagens periode er rykket videre), skal du først gå til Aktiviteter og "
        "bladre tilbage til den ønskede periode med '‹', og derefter klikke på Lønafregning."
    ))

    heading(doc, "Eksportér CSV", 2, "12.4")
    body(doc, (
        "Knappen '💾 Eksportér CSV' åbner en mappevælger (samme mønster som Lønkørsels "
        "'Kør løn') og gemmer én CSV-fil for hele den viste periode med data for alle "
        "medarbejdere – ikke topsummeringen."
    ))
    bullet(doc, "Naviger til den ønskede periode (se afsnit 12.3).")
    bullet(doc, "Klik '💾 Eksportér CSV'.")
    bullet(doc, "Vælg en mappe (forslået: din Downloads-mappe), eller klik 'Gennemse' for at vælge en anden.")
    bullet(doc, "Klik 'Eksportér'.")
    body(doc, (
        "CSV-filen (semikolon-separeret) indeholder kolonnerne Dato, Lønnummer, Normal timer, "
        "Overtid 1 time før, Overtid 1-3 timer efter, Øvrig overtid, Total tid, Total i kr., "
        "Vognnummer og Beløb – én række pr. dag pr. medarbejder, efterfulgt af en "
        "'Total løn for [navn]'-række."
    ))
    note_box(doc,
        "Eksport kræver at den viste periode er låst (der er kørt løn for den under "
        "Lønkørsel) – ellers vises en fejlbesked, og der eksporteres ikke. Administratorer "
        "er undtaget og kan altid eksportere, uanset periodens status. Selve eksporten LÅSER "
        "ikke perioden – det sker kun ved 'Kør løn' under Lønkørsel.",
        "VIGTIGT"
    )

    doc.save(OUT / "Brugervejledning.docx")
    print("Brugervejledning.docx gemt.")


if __name__ == "__main__":
    build_teknisk()
    build_bruger()
    print("Begge dokumenter er gemt i docs/-mappen.")
